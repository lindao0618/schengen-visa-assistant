from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "storage" / "templates" / "us-visa" / "interview-brief-template.docx"

TEXT = RGBColor(30, 41, 59)
MUTED = RGBColor(100, 116, 139)
ACCENT = RGBColor(14, 116, 144)
WARN = RGBColor(180, 83, 9)

PROCESS_STEPS = [
    (
        "1. 排队入场",
        "先去左边柜台 check in，工作人员会在你的 DS-160 Confirmation 上写预约时间，再让你去右边排队入场。",
    ),
    (
        "2. 入场安检",
        "可以带小包，外套和电子设备放在一起过安检，资料放到另一个篮子里即可。",
    ),
    (
        "3. 正式办理",
        "进入另一栋大楼后，出示护照和 DS-160 Confirmation，工作人员会给你号码，然后上楼正式办理。",
    ),
    (
        "4. 递交资料",
        "按照号码到对应窗口，递交护照原件、DS-160 Confirmation 和 BRP 原件。",
    ),
    (
        "5. 按指纹",
        "统一到左边窗口排队，扫 DS-160 Confirmation 上的条形码，然后按左四右四和两个大拇指指纹。",
    ),
    (
        "6. 面试",
        "递签后一周左右通常会收到两封邮件：第一封提示护照从大使馆发出，第二封提示护照抵达自取点。收到第二封后即可前往领取。",
    ),
]

PRACTICAL_NOTES = [
    "建议至少提前 30 分钟到场，先确认自己排的是 check in 队，再去安检，避免排错队伍。",
    "官方安检要求是只带面试必需物品；大包通常不允许带入，而且现场没有寄存服务。",
    "官方说明手机、电子书和平板通常可以携带，但不要带带键盘的平板或笔记本电脑。",
    "为减少安检卡住的风险，实操上也不要携带剪刀、刀具、指甲锉、玻璃瓶、打火机等尖锐或易碎物品。",
    "实操上也建议不要带皮带、手表和难脱的厚外套，安检时会更省时间。",
    "如果照片现场被要求重拍，通常可以在馆内 photo booth 处理，但最好还是提前准备好合规照片。",
    "材料尽量提前分好层次，护照、DS-160 Confirmation、BRP 放在最容易拿的位置。",
]


def set_run_style(run, *, size=11, bold=False, color=TEXT):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Cambria"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Cambria")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_paragraph(document: Document, text: str = "", *, size=11, bold=False, color=TEXT, after=4, line=1.35):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if text:
        run = paragraph.add_run(text)
        set_run_style(run, size=size, bold=bold, color=color)
    return paragraph


def set_cell_text(cell, text, *, bold=False, color=TEXT, bg=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_style(run, size=10.5, bold=bold, color=color)
    if bg:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), bg)


def build_template():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Cambria"
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Cambria")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria")
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal_style.font.size = Pt(10.5)

    add_paragraph(document, "美签递签与面试必看", size=18, bold=True, after=2)
    add_paragraph(
        document,
        "前半部分是递签流程与面试提醒，中间问答区会按申请人的个人信息做定制整理，最后保留取件与邮寄提醒。",
        size=10,
        color=MUTED,
        after=10,
    )

    add_paragraph(document, "签证中心位置", size=15, bold=True, color=ACCENT, after=6, line=1.15)
    location = document.add_table(rows=4, cols=2)
    set_cell_text(location.rows[0].cells[0], "Consular Section Location", bold=True, bg="E0F2FE")
    set_cell_text(location.rows[0].cells[1], "US Embassy London", bg="E0F2FE")
    set_cell_text(location.rows[1].cells[0], "City", bold=True)
    set_cell_text(location.rows[1].cells[1], "London")
    set_cell_text(location.rows[2].cells[0], "Postcode", bold=True)
    set_cell_text(location.rows[2].cells[1], "SW11 7US")
    set_cell_text(location.rows[3].cells[0], "Address", bold=True)
    set_cell_text(location.rows[3].cells[1], "33 Nine Elms Lane")

    add_paragraph(document, "", after=2)
    add_paragraph(document, "递签流程", size=15, bold=True, color=ACCENT, after=6, line=1.15)
    for title, body in PROCESS_STEPS:
        add_paragraph(document, title, size=12.5, bold=True, color=TEXT, after=2, line=1.15)
        add_paragraph(document, body, size=10.5, color=TEXT, after=6)

    add_paragraph(document, "面试当天注意事项", size=15, bold=True, color=ACCENT, after=6, line=1.15)
    for note in PRACTICAL_NOTES:
        add_paragraph(document, f"• {note}", size=10.5, color=TEXT, after=4)

    add_paragraph(document, "查询网址", size=12.5, bold=True, color=TEXT, after=4)
    add_paragraph(document, "https://ais.usvisa-info.com/en-gb/niv/users/sign_in", size=10.5, color=ACCENT, after=2)
    add_paragraph(document, "账号密码就是之前支付签证费用的账号密码。", size=10.5, color=TEXT, after=10)

    add_paragraph(document, "1.为什么去美国? Why are you going to the United States?", size=12.5, bold=True, after=2)
    add_paragraph(document, "去旅行", size=10.5, after=2)
    add_paragraph(document, "Go for a trip.", size=10.5, color=MUTED, after=8)

    add_paragraph(document, "伦敦自取点与取件提醒", size=15, bold=True, color=ACCENT, after=6, line=1.15)
    add_paragraph(document, "下面的信息固定保留，不会被中间的定制问答区覆盖。面试通过后优先按这里的取件方式提醒客户。", size=10.5, color=TEXT, after=6)

    free_pickup = document.add_table(rows=3, cols=2)
    set_cell_text(free_pickup.rows[0].cells[0], "免费自取点", bold=True, bg="E0F2FE")
    set_cell_text(free_pickup.rows[0].cells[1], "地址", bold=True, bg="E0F2FE")
    set_cell_text(free_pickup.rows[1].cells[0], "伦敦")
    set_cell_text(free_pickup.rows[1].cells[1], "Mail Boxes Etc. Holborn – 19 Bury Pl, London, WC1A 2JB")
    set_cell_text(free_pickup.rows[2].cells[0], "贝尔法斯特")
    set_cell_text(free_pickup.rows[2].cells[1], "Mail Boxes Etc. Belfast – Business Park, Belfast, BT6 9H")

    add_paragraph(document, "", after=2)
    add_paragraph(document, "曼城免费自取点停止服务", size=12.5, bold=True, color=WARN, after=2, line=1.15)
    add_paragraph(
        document,
        "所有之前选择曼城自取点的，官方会自动送到伦敦自取点；如果不想去伦敦拿，就只能选择快递到家或者付费的曼城自取点。",
        size=10.5,
        color=TEXT,
        after=8,
    )

    add_paragraph(document, "邮寄与收费自取", size=12.5, bold=True, color=ACCENT, after=2, line=1.15)
    courier = document.add_table(rows=4, cols=2)
    set_cell_text(courier.rows[0].cells[0], "项目", bold=True, bg="DBEAFE")
    set_cell_text(courier.rows[0].cells[1], "说明", bold=True, bg="DBEAFE")
    set_cell_text(courier.rows[1].cells[0], "快递到家")
    set_cell_text(courier.rows[1].cells[1], "美签快递到家邮费：30 镑")
    set_cell_text(courier.rows[2].cells[0], "伦敦收费自取")
    set_cell_text(courier.rows[2].cells[1], "Mail Boxes Etc. Angel – 8 Duncan Street, London, N1 8BW")
    set_cell_text(courier.rows[3].cells[0], "曼城收费自取")
    set_cell_text(courier.rows[3].cells[1], "Mail Boxes Etc. Manchester – 19 Lever Street, Manchester, M1 1AN")

    add_paragraph(document, "", after=2)
    add_paragraph(document, "收费自取点需要在递签之前先支付 23 镑服务费。", size=10.5, color=TEXT, after=2)

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(TEMPLATE_PATH))
    print(str(TEMPLATE_PATH))


if __name__ == "__main__":
    build_template()
