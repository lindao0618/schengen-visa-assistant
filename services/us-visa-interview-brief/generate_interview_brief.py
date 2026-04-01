import argparse
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


START_MARKERS = [
    "1.为什么去美国",
    "why are you going to the united states",
]

END_MARKERS = [
    "伦敦自取点",
]

ACCENT = RGBColor(14, 116, 144)
TEXT = RGBColor(30, 41, 59)
MUTED = RGBColor(100, 116, 139)
WARNING = RGBColor(180, 83, 9)


def normalize_text(value: str) -> str:
    return "".join(str(value or "").strip().lower().split())


def matches_any(paragraph_text: str, markers: list[str]) -> bool:
    normalized = normalize_text(paragraph_text)
    return any(normalize_text(marker) in normalized for marker in markers)


def find_anchor_indexes(document: Document) -> tuple[int, int]:
    start_index = -1
    end_index = -1

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text or ""
        if start_index < 0 and matches_any(text, START_MARKERS):
            start_index = index
            continue
        if start_index >= 0 and matches_any(text, END_MARKERS):
            end_index = index
            break

    if start_index < 0:
        raise RuntimeError("Could not find the start of the editable interview section.")
    if end_index < 0 or end_index <= start_index:
        raise RuntimeError("Could not find the end anchor of the editable interview section.")

    return start_index, end_index


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.3):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_paragraph_border(paragraph, color="E2E8F0"):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    border = paragraph_properties.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        paragraph_properties.append(border)

    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = border.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            border.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), color)


def set_run_font(run, *, size=11, bold=False, color=TEXT):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Cambria"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Cambria")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_labelled_line(paragraph, label: str, value: str, *, label_color=ACCENT, value_color=TEXT):
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=10.5, bold=True, color=label_color)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5, color=value_color)


def insert_heading(anchor, title: str, *, size=15, color=ACCENT):
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=10, after=5, line=1.1)
    run = paragraph.add_run(title)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def insert_body(anchor, text: str, *, size=10.5, color=TEXT, before=0, after=0):
    paragraph = anchor.insert_paragraph_before()
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.35)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color)
    return paragraph


def render_standard_block(anchor, block: dict):
    insert_heading(anchor, block.get("questionCn", ""), size=13)
    if block.get("questionEn"):
        subtitle = anchor.insert_paragraph_before()
        set_paragraph_spacing(subtitle, before=0, after=4, line=1.15)
        run = subtitle.add_run(block["questionEn"])
        set_run_font(run, size=9.8, color=MUTED)

    answer_box = anchor.insert_paragraph_before()
    set_paragraph_spacing(answer_box, before=2, after=6, line=1.4)
    set_paragraph_border(answer_box)
    add_labelled_line(answer_box, "中文：", block.get("answerCn", ""))

    en_line = anchor.insert_paragraph_before()
    set_paragraph_spacing(en_line, before=0, after=8, line=1.4)
    set_paragraph_border(en_line)
    add_labelled_line(en_line, "英文：", block.get("answerEn", ""), value_color=MUTED)

    if block.get("note"):
        note = insert_body(anchor, block["note"], size=9.5, color=WARNING, after=8)
        note.paragraph_format.left_indent = Pt(6)


def render_options_block(anchor, block: dict):
    render_standard_block(
        anchor,
        {
            "questionCn": block.get("questionCn", ""),
            "questionEn": block.get("questionEn"),
            "answerCn": block.get("answerCn", ""),
            "answerEn": block.get("answerEn", ""),
        },
    )

    if block.get("note"):
        note = insert_body(anchor, block["note"], size=9.5, color=WARNING, after=6)
        note.paragraph_format.left_indent = Pt(6)

    for option in block.get("options", []):
        option_title = anchor.insert_paragraph_before()
        set_paragraph_spacing(option_title, before=0, after=0, line=1.1)
        label_run = option_title.add_run(f"{option.get('label', '').strip()}：")
        set_run_font(label_run, size=10, bold=True, color=ACCENT)

        option_body = anchor.insert_paragraph_before()
        set_paragraph_spacing(option_body, before=0, after=6, line=1.35)
        set_paragraph_border(option_body, color="BAE6FD")
        body_run = option_body.add_run(option.get("text", ""))
        set_run_font(body_run, size=10.6, color=TEXT)


def render_section_title(anchor, block: dict):
    insert_heading(anchor, block.get("title", ""), size=14, color=RGBColor(194, 65, 12))
    if block.get("description"):
        description = insert_body(anchor, block["description"], size=10, color=RGBColor(120, 53, 15), after=8)
        description.paragraph_format.left_indent = Pt(4)


def render_pending_block(anchor, block: dict):
    insert_heading(anchor, block.get("questionCn", ""), size=13)
    if block.get("questionEn"):
        subtitle = anchor.insert_paragraph_before()
        set_paragraph_spacing(subtitle, before=0, after=4, line=1.15)
        run = subtitle.add_run(block["questionEn"])
        set_run_font(run, size=9.8, color=MUTED)

    cn_box = anchor.insert_paragraph_before()
    set_paragraph_spacing(cn_box, before=2, after=4, line=1.35)
    set_paragraph_border(cn_box, color="FCD34D")
    add_labelled_line(cn_box, "中文：", block.get("placeholderCn", ""), label_color=WARNING, value_color=WARNING)

    en_box = anchor.insert_paragraph_before()
    set_paragraph_spacing(en_box, before=0, after=8, line=1.35)
    set_paragraph_border(en_box, color="FCD34D")
    add_labelled_line(en_box, "英文：", block.get("placeholderEn", ""), label_color=WARNING, value_color=WARNING)


def render_hotel_block(anchor, block: dict):
    insert_heading(anchor, block.get("title", ""), size=13)
    hotel = anchor.insert_paragraph_before()
    set_paragraph_spacing(hotel, before=2, after=10, line=1.35)
    set_paragraph_border(hotel, color="BFDBFE")
    run = hotel.add_run(block.get("hotelName", ""))
    set_run_font(run, size=10.6, bold=True, color=RGBColor(3, 105, 161))


def replace_section(document: Document, blocks: list[dict]) -> None:
    start_index, end_index = find_anchor_indexes(document)
    anchor = document.paragraphs[end_index]
    removable = list(document.paragraphs[start_index:end_index])

    for paragraph in removable:
        remove_paragraph(paragraph)

    intro = anchor.insert_paragraph_before()
    set_paragraph_spacing(intro, before=8, after=10, line=1.2)
    intro_run = intro.add_run("美签常见问题")
    set_run_font(intro_run, size=16, bold=True, color=RGBColor(15, 23, 42))

    intro_desc = anchor.insert_paragraph_before()
    set_paragraph_spacing(intro_desc, before=0, after=10, line=1.35)
    desc_run = intro_desc.add_run("以下内容已根据申请人的个人信息定制整理；中文和英文都保留，方便面试前快速记忆。")
    set_run_font(desc_run, size=10, color=MUTED)

    for block in blocks:
      block_type = block.get("type")
      if block_type == "qa":
          render_standard_block(anchor, block)
      elif block_type == "qa-options":
          render_options_block(anchor, block)
      elif block_type == "section-title":
          render_section_title(anchor, block)
      elif block_type == "pending-qa":
          render_pending_block(anchor, block)
      elif block_type == "hotel":
          render_hotel_block(anchor, block)
      else:
          raise RuntimeError(f"Unsupported block type: {block_type}")


def try_docx2pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception:
        return False


def try_libreoffice(docx_path: Path, pdf_path: Path) -> bool:
    programs = ["soffice", "libreoffice"]
    out_dir = str(pdf_path.parent)

    for program in programs:
        try:
            subprocess.run(
                [program, "--headless", "--convert-to", "pdf", "--outdir", out_dir, str(docx_path)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )
            generated = docx_path.with_suffix(".pdf")
            if generated.exists() and generated != pdf_path:
                shutil.move(str(generated), str(pdf_path))
            return pdf_path.exists()
        except Exception:
            continue

    return False


def try_word_com(docx_path: Path, pdf_path: Path) -> bool:
    if os.name != "nt":
        return False

    word = None
    document = None
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(docx_path),
            ReadOnly=True,
            AddToRecentFiles=False,
            ConfirmConversions=False,
            NoEncodingDialog=True,
        )
        document.ExportAsFixedFormat(str(pdf_path), 17)
        return pdf_path.exists()
    except Exception:
        return False
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str | None]:
    if try_docx2pdf(docx_path, pdf_path):
        return True, None
    if try_libreoffice(docx_path, pdf_path):
        return True, None
    if try_word_com(docx_path, pdf_path):
        return True, None
    return False, "PDF conversion is unavailable on this machine."


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--payload", required=True)
    parser.add_argument("--output-dir", required=True)
    args = parser.parse_args()

    template_path = Path(args.template)
    payload_path = Path(args.payload)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    blocks = payload.get("blocks") or []
    if not isinstance(blocks, list) or not blocks:
        raise RuntimeError("Missing blocks in payload.")

    document = Document(str(template_path))
    replace_section(document, blocks)

    docx_name = payload.get("docxFilename") or "us_visa_interview_brief.docx"
    pdf_name = payload.get("pdfFilename") or "us_visa_interview_brief.pdf"
    docx_path = output_dir / docx_name
    pdf_path = output_dir / pdf_name

    document.save(str(docx_path))
    pdf_ok, pdf_warning = convert_to_pdf(docx_path, pdf_path)

    result = {
        "success": True,
        "docx_file": docx_path.name,
        "pdf_file": pdf_path.name if pdf_ok else None,
        "pdf_warning": pdf_warning,
    }
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(json.dumps({"success": False, "error": str(exc)}, ensure_ascii=False))
        sys.exit(1)
