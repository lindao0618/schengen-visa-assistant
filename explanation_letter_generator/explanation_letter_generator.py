"""
解释信文档生成器
用于创建Word文档和PDF转换
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging
from datetime import datetime
import os
from typing import Optional

try:
    from docx2pdf import convert
    PDF_CONVERSION_AVAILABLE = True
except ImportError:
    PDF_CONVERSION_AVAILABLE = False
    logging.warning("docx2pdf not available, PDF conversion will be disabled")

logger = logging.getLogger(__name__)

class ExplanationLetterGenerator:
    def __init__(self):
        self.document = None
    
    def create_explanation_letter_document(
        self,
        content: str,
        chinese_name: str,
        english_name: str,
        passport_number: str,
        output_path: str,
        english_content: Optional[str] = None,
    ):
        """
        创建解释信Word文档，可选包含英文版本
        """
        try:
            # 创建新文档
            self.document = Document()

            # 设置页面边距
            sections = self.document.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

            # 添加标题
            title_paragraph = self.document.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run("解释信")
            title_run.font.name = "Microsoft YaHei"
            title_run.font.size = Pt(18)
            title_run.bold = True

            # 添加空行
            self.document.add_paragraph()

            # 添加中文解释信内容
            content_lines = content.split('\n')
            for line in content_lines:
                if line.strip():  # 跳过空行
                    paragraph = self.document.add_paragraph(line.strip())
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
                    
                    # 设置段落格式
                    for run in paragraph.runs:
                        run.font.name = "Microsoft YaHei"
                        run.font.size = Pt(12)
                    
                    # 设置段落间距
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15
                else:
                    # 添加空行
                    self.document.add_paragraph()
            
            # 添加签名区域（姓名+日期，左对齐）
            self.document.add_paragraph()
            self.document.add_paragraph()
            name_para = self.document.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_run = name_para.add_run(f"{chinese_name}（{english_name}）")
            name_run.font.name = "Microsoft YaHei"
            name_run.font.size = Pt(12)
            date_para = self.document.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            date_run = date_para.add_run(datetime.now().strftime('%Y年%m月%d日'))
            date_run.font.name = "Microsoft YaHei"
            date_run.font.size = Pt(12)

            # 如有英文版本，添加分页及英文内容
            if english_content:
                from docx.enum.text import WD_BREAK
                self.document.add_paragraph()
                page_break_para = self.document.add_paragraph()
                run = page_break_para.add_run()
                run.add_break(WD_BREAK.PAGE)

                en_title = self.document.add_paragraph()
                en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                en_title_run = en_title.add_run("Explanation Letter (English)")
                en_title_run.font.name = "Times New Roman"
                en_title_run.font.size = Pt(18)
                en_title_run.bold = True
                self.document.add_paragraph()

                for line in english_content.split('\n'):
                    if line.strip():
                        para = self.document.add_paragraph(line.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for r in para.runs:
                            r.font.name = "Times New Roman"
                            r.font.size = Pt(12)
                        para.paragraph_format.space_after = Pt(6)
                        para.paragraph_format.line_spacing = 1.15
                    else:
                        self.document.add_paragraph()

                self.document.add_paragraph()
                self.document.add_paragraph()
                sig_en = self.document.add_paragraph()
                sig_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sig_en_run = sig_en.add_run(english_name)
                sig_en_run.font.name = "Times New Roman"
                sig_en_run.font.size = Pt(12)
                date_en = self.document.add_paragraph()
                date_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
                date_en_run = date_en.add_run(datetime.now().strftime('%Y-%m-%d'))
                date_en_run.font.name = "Times New Roman"
                date_en_run.font.size = Pt(12)

            # 保存文档
            self.document.save(output_path)
            logger.info(f"解释信Word文档已保存：{output_path}")

        except Exception as e:
            logger.error(f"创建解释信文档失败: {str(e)}")
            raise

    def create_english_document(
        self,
        content: str,
        english_name: str,
        passport_number: str,
        output_path: str,
    ):
        """创建纯英文解释信 Word 文档"""
        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_p.add_run("Cover Letter")
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)
            run.bold = True
            doc.add_paragraph()

            for line in content.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for r in para.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(12)
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.line_spacing = 1.15
                else:
                    doc.add_paragraph()

            date_p = doc.add_paragraph()
            date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            date_run = date_p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            date_run.font.name = "Times New Roman"
            date_run.font.size = Pt(12)

            doc.save(output_path)
            logger.info(f"英文解释信Word已保存：{output_path}")
        except Exception as e:
            logger.error(f"创建英文解释信失败: {str(e)}")
            raise

    def convert_to_pdf(self, docx_path: str, pdf_path: str):
        """
        将Word文档转换为PDF（docx2pdf 优先，失败则尝试 LibreOffice）
        """
        import subprocess
        import time

        def try_docx2pdf():
            if not PDF_CONVERSION_AVAILABLE:
                return False
            try:
                convert(docx_path, pdf_path)
                return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 500
            except Exception:
                return False

        def try_libreoffice():
            import shutil
            abs_docx = os.path.abspath(docx_path)
            out_dir = os.path.dirname(abs_docx)
            for prog in ["soffice", "libreoffice"]:
                try:
                    subprocess.run(
                        [prog, "--headless", "--convert-to", "pdf", "--outdir", out_dir, abs_docx],
                        capture_output=True, timeout=30
                    )
                    base = os.path.splitext(os.path.basename(docx_path))[0]
                    expected_pdf = os.path.join(out_dir, base + ".pdf")
                    if os.path.exists(expected_pdf) and os.path.getsize(expected_pdf) > 500:
                        abs_pdf = os.path.abspath(pdf_path)
                        if os.path.abspath(expected_pdf) != abs_pdf:
                            shutil.copy2(expected_pdf, abs_pdf)
                            try:
                                os.remove(expected_pdf)
                            except OSError:
                                pass
                        return True
                except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
                    continue
            return False

        try:
            if try_docx2pdf():
                logger.info(f"PDF转换成功：{pdf_path}")
                return
            time.sleep(1)
            if try_docx2pdf():
                logger.info(f"PDF转换成功（重试）：{pdf_path}")
                return
            if try_libreoffice():
                logger.info(f"PDF转换成功（LibreOffice）：{pdf_path}")
                return
            logger.warning("PDF转换失败，将生成占位文件")
            with open(pdf_path, 'w', encoding='utf-8') as f:
                f.write("PDF转换不可用，请下载 Word 文档。")
        except Exception as e:
            logger.error(f"PDF转换失败: {str(e)}")
            with open(pdf_path, 'w', encoding='utf-8') as f:
                f.write(f"PDF转换失败：{str(e)}\n请使用Word文档")
    
    def create_template_document(self, template_path: str):
        """
        创建解释信模板文档（用于测试）
        """
        try:
            document = Document()
            
            # 设置页面边距
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # 添加标题
            title_paragraph = document.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run("解释信模板")
            title_run.font.name = "Microsoft YaHei"
            title_run.font.size = Pt(18)
            title_run.bold = True
            
            # 添加模板说明
            document.add_paragraph()
            content = """
致签证中心：

我是[申请人姓名]，护照号码：[护照号码]，现就职于[工作单位]。

我计划于[出行日期]前往[目的国家]，申请[签证类型]签证。在准备签证材料过程中，发现以下问题需要向贵处说明：

[问题详细说明]

我理解这一情况可能会对签证审核造成疑虑，特此提供详细说明，并承诺所有信息真实有效。如需进一步材料或说明，我将积极配合提供。

恳请贵处理解我的实际情况，批准我的签证申请。

此致
敬礼！

申请人：[申请人姓名]
护照号码：[护照号码]
日期：[日期]
"""
            
            content_lines = content.strip().split('\n')
            for line in content_lines:
                if line.strip():
                    paragraph = document.add_paragraph(line.strip())
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    for run in paragraph.runs:
                        run.font.name = "Microsoft YaHei"
                        run.font.size = Pt(12)
                    
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15
                else:
                    document.add_paragraph()
            
            # 保存模板
            document.save(template_path)
            logger.info(f"解释信模板已创建：{template_path}")
            
        except Exception as e:
            logger.error(f"创建解释信模板失败: {str(e)}")
            raise