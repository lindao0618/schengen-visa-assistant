from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from fastapi.responses import JSONResponse
from datetime import datetime, timedelta
from prompt_builder import build_prompt
from deepseek_api import call_deepseek
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Inches, Pt
import shutil
import os
import re
import base64
import logging
import subprocess
from pathlib import Path
from copy import deepcopy

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except Exception:
    colors = None
    A4 = None
    landscape = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    mm = None
    pdfmetrics = None
    UnicodeCIDFont = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TRIP_GENERATOR_RUNTIME_DIR = PROJECT_ROOT / "temp" / "trip_generator"
TRIP_GENERATOR_OUTPUT_DIR = TRIP_GENERATOR_RUNTIME_DIR / "output"
TEMPLATE_DOCX_PATH = Path(__file__).resolve().parent / "template.docx"

app = FastAPI()

@app.get("/health")
def health():
    """健康检查：用于确认行程单服务是否已启动"""
    return {"status": "ok", "service": "trip_generator"}

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TABLE_COLUMN_WIDTHS_MM = [14, 24, 32, 84, 108]


def get_libreoffice_candidates() -> list[str]:
    candidates: list[str] = []

    env_bin = os.environ.get("LIBREOFFICE_BIN", "").strip()
    if env_bin:
        candidates.append(env_bin)

    for name in ("soffice", "libreoffice"):
        resolved = shutil.which(name)
        if resolved:
            candidates.append(resolved)

    if os.name == "nt":
        candidates.extend(
            [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
        )

    seen: set[str] = set()
    unique_candidates: list[str] = []
    for candidate in candidates:
        if candidate and candidate not in seen:
            seen.add(candidate)
            unique_candidates.append(candidate)
    return unique_candidates


def convert_docx_to_pdf(output_docx: str, output_pdf: str):
    if docx2pdf_convert is not None:
        try:
            logger.info("Converting Word to PDF via docx2pdf...")
            docx2pdf_convert(output_docx, output_pdf)
            if os.path.exists(output_pdf):
                logger.info(f"PDF saved as: {output_pdf}")
                return
        except Exception as exc:
            logger.warning(f"docx2pdf failed ({exc}), trying LibreOffice...")
    else:
        logger.info("docx2pdf unavailable, trying LibreOffice...")

    abs_docx = os.path.abspath(output_docx)
    out_dir = os.path.dirname(abs_docx)
    generated_pdf = os.path.splitext(abs_docx)[0] + ".pdf"
    conversion_errors: list[str] = []

    for prog in get_libreoffice_candidates():
        try:
            logger.info(f"Trying LibreOffice converter: {prog}")
            result = subprocess.run(
                [prog, "--headless", "--convert-to", "pdf", "--outdir", out_dir, abs_docx],
                capture_output=True,
                text=True,
                timeout=120,
                cwd=os.getcwd(),
            )
            logger.info(
                "LibreOffice conversion result",
                extra={
                    "returncode": result.returncode,
                    "stdout": (result.stdout or "").strip(),
                    "stderr": (result.stderr or "").strip(),
                },
            )
            if result.returncode == 0 and os.path.exists(generated_pdf):
                if os.path.abspath(generated_pdf) != os.path.abspath(output_pdf):
                    shutil.move(generated_pdf, output_pdf)
                logger.info("PDF created via LibreOffice")
                return
            conversion_errors.append(
                f"{prog}: code={result.returncode}, stdout={result.stdout.strip()}, stderr={result.stderr.strip()}"
            )
        except FileNotFoundError:
            conversion_errors.append(f"{prog}: not found")
        except subprocess.TimeoutExpired:
            conversion_errors.append(f"{prog}: timeout")
        except Exception as exc:
            conversion_errors.append(f"{prog}: {exc}")

    raise RuntimeError(
        "PDF 生成失败。当前未检测到可用的 Word/LibreOffice 转换器。"
        + (" 详情: " + " | ".join(conversion_errors) if conversion_errors else "")
    )


def create_pdf_directly(rows, output_pdf: str):
    if not all(
        [
            colors,
            A4,
            landscape,
            ParagraphStyle,
            getSampleStyleSheet,
            mm,
            pdfmetrics,
            UnicodeCIDFont,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        ]
    ):
        raise RuntimeError("缺少 reportlab，无法使用纯 Python 直接生成 PDF。")

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ItineraryTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=10,
    )
    header_style = ParagraphStyle(
        "ItineraryHeader",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=12,
        alignment=1,
        textColor=colors.white,
    )
    center_style = ParagraphStyle(
        "ItineraryCenterCell",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=9,
        leading=11,
        alignment=1,
    )
    left_style = ParagraphStyle(
        "ItineraryLeftCell",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=9,
        leading=11,
        alignment=0,
    )

    doc = SimpleDocTemplate(
        output_pdf,
        pagesize=landscape(A4),
        leftMargin=8 * mm,
        rightMargin=8 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )
    story = [Paragraph("Travel Itinerary", title_style), Spacer(1, 4 * mm)]
    table_data = [
        [
            Paragraph("Day", header_style),
            Paragraph("Date", header_style),
            Paragraph("Transport", header_style),
            Paragraph("Attractions", header_style),
            Paragraph("Hotel Information", header_style),
        ]
    ]

    for row in rows:
        rendered = []
        for index, value in enumerate(row):
            text = str(value or "").replace("\n", "<br/>")
            rendered.append(Paragraph(text, center_style if index in (0, 1, 2) else left_style))
        table_data.append(rendered)

    table = Table(table_data, colWidths=[width * mm for width in TABLE_COLUMN_WIDTHS_MM], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9ca3af")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ]
        )
    )
    story.append(table)
    doc.build(story)


def apply_table_layout(doc: Document, table):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row in table.rows:
        for index, width_mm in enumerate(TABLE_COLUMN_WIDTHS_MM):
            if index < len(row.cells):
                row.cells[index].width = Mm(width_mm)

    if len(table.rows) >= 1:
        table.rows[0].height = Mm(10)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    if len(table.rows) >= 2:
        table.rows[1].height = Mm(9)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for row in table.rows[2:]:
        row.height = Mm(18)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def apply_cell_style(cell, *, font_size, align, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05

        for run in paragraph.runs:
            run.font.size = font_size
            if bold:
                run.bold = True


def apply_table_text_style(table):
    for cell in table.rows[0].cells:
        apply_cell_style(cell, font_size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    for cell in table.rows[1].cells:
        apply_cell_style(cell, font_size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    for row in table.rows[2:]:
        for index, cell in enumerate(row.cells):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            apply_cell_style(cell, font_size=Pt(11), align=align)

def clear_cell_text_preserve_style(cell):
    while len(cell.paragraphs) > 1:
        cell._tc.remove(cell.paragraphs[-1]._element)

    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    while len(paragraph.runs) > 1:
        paragraph._element.remove(paragraph.runs[-1]._element)

    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = ""
    else:
        run = paragraph.add_run("")

    return run

def create_word_from_template(rows, template_path: str, output_docx: str):
    """Create Word document by cloning the template's sample data row."""

    doc = Document(template_path)
    if not doc.tables:
        raise RuntimeError("Template document does not contain any table.")

    table = doc.tables[0]
    if len(table.rows) < 3:
        raise RuntimeError("Template table must contain title row, header row, and one sample data row.")
    
    print(f"Template table has {len(table.rows)} rows")
    print(f"Data to insert: {len(rows)} rows")
    print(f"Data content: {rows}")

    template_row_xml = deepcopy(table.rows[2]._tr)

    # 保留前2行（标题行和表头），清除模板里的示例数据行
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)
    
    print(f"After clearing, table has {len(table.rows)} rows")

    # 为每个数据行克隆模板数据行，保持模板本身的样式
    for i, row_data in enumerate(rows):
        print(f"Processing row {i}: {row_data}")
        table._tbl.append(deepcopy(template_row_xml))
        cells = table.rows[-1].cells

        for j, content in enumerate(row_data):
            if j >= len(cells):
                continue
            cell = cells[j]
            run = clear_cell_text_preserve_style(cell)
            run.text = str(content or "")
            print(f"  Cell {j}: '{content}' -> '{run.text}'")

    print(f"Final table has {len(table.rows)} rows")
    doc.save(output_docx)
    print(f"Document saved to {output_docx}")

def create_pdf_from_template(rows, template_path: str, output_pdf: str):
    """Create PDF using the template file"""
    import os
    from docx2pdf import convert
    
    # 先生成Word文件
    output_docx = output_pdf.replace('.pdf', '.docx')
    create_word_from_template(rows, template_path, output_docx)
    
    # 然后转换为PDF
    convert(output_docx, output_pdf)
    
    # 删除临时Word文件
    os.remove(output_docx)

class ItineraryRequest(BaseModel):
    country: str
    departure_city: str
    arrival_city: str
    start_date: str  # yyyy-mm-dd
    end_date: str    # yyyy-mm-dd
    hotel_name: str
    hotel_address: str
    hotel_phone: str

class ItineraryResponse(BaseModel):
    pdf_base64: str
    analysis: str

def parse_attractions(response):
    """Parse AI response to extract dates and attractions"""
    # Split response by lines and remove empty lines
    lines = [line.strip() for line in response.split('\n') if line.strip()]
    
    daily_attractions = {}
    for line in lines:
        # Look for lines containing date and attractions (format: "YYYY.MM.DD: Attraction1, Attraction2")
        if ':' in line:
            date_str, attractions = line.split(':', 1)
            date_str = date_str.strip()
            attractions = attractions.strip()
            
            # Convert YYYY.MM.DD format to DD/MM/YYYY format
            if len(date_str) == 10 and date_str[4] == '.' and date_str[7] == '.':
                year = date_str[0:4]
                month = date_str[5:7]
                day = date_str[8:10]
                date_str = f"{day}/{month}/{year}"
            
            # Special handling for Palace of Versailles
            if "Palace of Versailles" in attractions or "Château de Versailles" in attractions:
                print(f"\n🏰 Found Palace of Versailles {date_str} - Setting as only attraction for the day")
                daily_attractions[date_str] = "Palace of Versailles"
            else:
                # Process comma-separated attractions
                attractions_list = [attr.strip() for attr in attractions.split(',')]
                daily_attractions[date_str] = ', '.join(attractions_list)
    
    return daily_attractions

def analyze_itinerary(rows):
    """Analyze the itinerary using Deepseek"""
    # Convert rows to readable format for analysis
    itinerary_text = "Generated Itinerary:\n\n"
    for row in rows:
        day, date, transport, spots, hotel = row
        itinerary_text += f"Day {day} ({date}):\n"
        itinerary_text += f"City: {transport}\n"
        if spots:
            itinerary_text += f"Attractions: {spots}\n"
        if hotel:
            itinerary_text += f"Hotel Information:\n{hotel}\n"
        itinerary_text += "-------------------\n"

    prompt = f"""Please analyze this travel itinerary based on French visa center (TLScontact/VFS Global) official requirements and visa officer practices. Consider:

1. Official Visa Requirements:
   - Are all dates listed consecutively?
   - Are all cities and transportation methods clearly indicated?
   - Are attractions properly planned for each day?（if there is only one attraction, it should be Palace of Versailles,this is suitable for visa officer）
   - Does it align with hotel bookings (check-in/check-out)?

2. Daily Schedule Credibility:
   - Are there 2-3 attractions per day (minimum requirement)?（if there is only one attraction, it should be Palace of Versailles,this is suitable for visa officer）
   - Is the first day marked with check-in and last day with check-out?
   - Are attractions logically varied between days (not repetitive)?
   - Is the number of attractions reasonable (not exceeding 4 per day)?
   - Do hotel locations match the cities visited?

3. Schedule Practicality:
   - Is the route between attractions efficient?
   - Is there enough time between attractions?
   - Are attractions logically grouped by location?
   - Are there any scheduling conflicts or impractical arrangements?

4. Overall Credibility:
   - Does the itinerary appear authentic and well-planned?
   - Are there emotional touchpoints (e.g., evening walks, relaxation time)?
   - Is there enough variety in activities?
   - Is the format clear and easy for visa officers to review?

Itinerary to analyze:
{itinerary_text}

请简要分析该行程是否符合签证基本要求和标准。"""

    response = call_deepseek(prompt, lang="zh")
    return response

@app.post("/generate-itinerary", response_model=ItineraryResponse)
def generate_itinerary(req: ItineraryRequest):
    try:
        # Date processing
        start_date = datetime.strptime(req.start_date, "%Y-%m-%d")
        end_date = datetime.strptime(req.end_date, "%Y-%m-%d")
        days = (end_date - start_date).days + 1

        print("\n=== Generated Daily Attractions ===")
        print(f"Destination: from {req.departure_city} to {req.arrival_city}")
        print(f"Duration: {days} days ({req.start_date} to {req.end_date})")
        print("=====================================\n")

        # Generate one prompt for the entire itinerary
        prompt = build_prompt(
            city=req.arrival_city,
            start=start_date.strftime("%Y.%m.%d"),
            end=end_date.strftime("%Y.%m.%d"),
            hotel={
                "名称": req.hotel_name,
                "地址": req.hotel_address,
                "电话": req.hotel_phone
            },
            lang="zh"
        )
        
        response = call_deepseek(prompt, lang="zh")
        print("\nDeepseek response:")
        print(response)
        
        daily_attractions = parse_attractions(response)
        daily_spots = {}
        
        # Process attractions for all days except first and last
        for i in range(1, days - 1):
            current_date = start_date + timedelta(days=i)
            display_date_str = current_date.strftime("%d/%m/%Y")
            
            if daily_attractions:
                attractions = daily_attractions.get(display_date_str, "")
                if attractions:
                    print(f"\nDay {i} ({display_date_str}):")
                    print(f"Attractions: {attractions}")
                    print("----------------------------------------")
                    daily_spots[display_date_str] = attractions

        rows = []
        for i in range(days):
            current_date = start_date + timedelta(days=i)
            date_str = current_date.strftime("%d/%m/%Y")
            
            # Handle transportation and attractions based on day type
            if i == 0:  # First day (arrival)
                transport = f"{req.departure_city} ➔ {req.arrival_city}"
                spots = "Arrival Day - Hotel Check-in"
                hotel_info = f"Hotel: {req.hotel_name}\nAddress: {req.hotel_address}\nPhone: {req.hotel_phone}"
            elif i == days - 1:  # Last day (departure)
                transport = f"{req.arrival_city} ➔ {req.departure_city}"
                spots = "Departure Day - Hotel Check-out"
                hotel_info = ""
            else:  # Regular sightseeing day
                transport = f"{req.arrival_city}"  # Just use arrival city name for middle days
                spots = daily_spots.get(date_str, "")  # Empty string if no attractions found
                hotel_info = f"Hotel: {req.hotel_name}\nAddress: {req.hotel_address}\nPhone: {req.hotel_phone}"

            rows.append([str(i + 1), date_str, transport, spots, hotel_info])

        # Write generated files outside app/ to avoid Next dev hot-reload loops.
        TRIP_GENERATOR_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_docx = str(TRIP_GENERATOR_OUTPUT_DIR / f"itinerary_{timestamp}.docx")
        output_pdf = str(TRIP_GENERATOR_OUTPUT_DIR / f"itinerary_{timestamp}.pdf")

        # Get itinerary analysis
        logger.info("Analyzing itinerary")
        analysis = analyze_itinerary(rows)
        
        logger.info("Creating Word document from template")
        create_word_from_template(rows, str(TEMPLATE_DOCX_PATH), output_docx)
        logger.info(f"Word document saved as: {output_docx}")
        convert_docx_to_pdf(output_docx, output_pdf)

        logger.info("Reading generated file")
        with open(output_pdf, "rb") as pdf_file:
            pdf_content = pdf_file.read()
            logger.info(f"File size: {len(pdf_content)} bytes")
            pdf_base64 = base64.b64encode(pdf_content).decode()
            logger.info(f"Base64 string length: {len(pdf_base64)}")

        # Return both PDF and analysis
        response_data = {
            "pdf_base64": pdf_base64,
            "analysis": analysis
        }
        
        logger.info("Sending response")
        return JSONResponse(content=response_data)
        
    except Exception as e:
        logger.error(f"Error in generate_itinerary: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

if __name__ == "__main__":
    import uvicorn
    try:
        print("Starting trip generator service on port 8002...")
        uvicorn.run(app, host="0.0.0.0", port=8002, log_level="info")
    except KeyboardInterrupt:
        print("Service stopped by user")
    except Exception as e:
        print(f"Service error: {e}")
        # Keep the service running even if there's an error
        import time
        while True:
            time.sleep(1)
