from docx import Document

def check_template_structure():
    """检查模板文件的结构"""
    doc = Document("template.docx")
    
    print("=== 模板文件结构检查 ===")
    print(f"文档中的表格数量: {len(doc.tables)}")
    
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"第一个表格的行数: {len(table.rows)}")
        print(f"第一个表格的列数: {len(table.columns)}")
        
        print("\n=== 表格内容 ===")
        for i, row in enumerate(table.rows):
            print(f"第 {i+1} 行:")
            for j, cell in enumerate(row.cells):
                text = cell.text.strip()
                print(f"  列 {j+1}: '{text}'")
                if len(cell.paragraphs) > 0:
                    para = cell.paragraphs[0]
                    print(f"    段落对齐方式: {para.alignment}")
                    if len(para.runs) > 0:
                        run = para.runs[0]
                        print(f"    字体大小: {run.font.size}")
                        print(f"    字体名称: {run.font.name}")
            print()

if __name__ == "__main__":
    check_template_structure() 