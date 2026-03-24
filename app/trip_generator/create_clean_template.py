#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建一个完全干净的模板，不包含任何硬编码数据
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_clean_template():
    """创建一个完全干净的模板文件"""
    
    # 读取原始模板以保持格式设置
    if os.path.exists("template.docx"):
        doc = Document("template.docx")
        # 保持原始页面设置
    else:
        doc = Document()
        
    # 清除所有内容
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Trip Itinerary')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_before = Pt(36)
    title_para.space_after = Pt(24)
    
    # 创建表格（只需要2行：表头+数据行）
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置列宽
    table.columns[0].width = Inches(0.6)   # Day
    table.columns[1].width = Inches(1.2)   # Date  
    table.columns[2].width = Inches(1.5)   # City
    table.columns[3].width = Inches(2.5)   # Touring spots
    table.columns[4].width = Inches(2.0)   # Accommodation
    
    # 设置表头（第一行）
    header_row = table.rows[0]
    headers = ['Day', 'Date', 'City', 'Touring spots', 'Accommodation']
    
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
    
    # 设置示例数据行（第二行）- 保持空白，只设置格式
    data_row = table.rows[1]
    data_values = ['1', '01/08/2025', 'CITY', 'Sample activities', 'Sample accommodation']
    
    for i, sample_text in enumerate(data_values):
        cell = data_row.cells[i]
        cell.text = sample_text
        
        para = cell.paragraphs[0]
        if i == 4:  # Accommodation列左对齐
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        run = para.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    # 保存模板
    doc.save("template_clean.docx")
    print("✅ 创建了干净的模板: template_clean.docx")

if __name__ == "__main__":
    import os
    create_clean_template()