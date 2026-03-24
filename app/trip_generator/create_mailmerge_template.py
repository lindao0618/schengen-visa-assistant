#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自动创建Mail Merge模板工具
将原始模板转换为带有docxtpl变量的Mail Merge模板
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def create_mailmerge_template():
    """创建Mail Merge模板"""
    
    print("📝 开始创建Mail Merge模板...")
    
    # 如果存在原模板，则基于原模板创建，保持边距设置
    if os.path.exists("template.docx"):
        print("   基于原模板创建，保持边距设置...")
        doc = Document("template.docx")
        
        # 清除现有内容
        for paragraph in doc.paragraphs:
            paragraph.clear()
        
        # 清除现有表格
        for table in doc.tables:
            table._element.getparent().remove(table._element)
    else:
        # 创建新文档
        doc = Document()
        
        # 设置页面边距 (更宽的边距)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.2)
            section.bottom_margin = Inches(1.2)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
    
    # 标题段落 - 设置最大边距
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('{{trip_title}}')
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置标题段落的上边距为最大
    title_para.paragraph_format.space_before = Pt(36)  # 最大上边距
    title_para.paragraph_format.space_after = Pt(24)   # 下边距也增大
    
    # 基本信息段落
    info_para = doc.add_paragraph()
    info_run = info_para.add_run('旅行日期：{{start_date}} 至 {{end_date}} ({{total_days}}天)')
    info_run.font.name = '微软雅黑'
    info_run.font.size = Pt(12)
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    doc.add_paragraph()
    
    # 创建表格 (4列) - 设置为页面居中
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格在页面中居中
    
    # 设置列宽
    table.columns[0].width = Inches(1.2)  # 日期
    table.columns[1].width = Inches(3.5)  # 行程安排
    table.columns[2].width = Inches(2.0)  # 住宿
    table.columns[3].width = Inches(1.0)  # 备注
    
    # 表头
    header_cells = table.rows[0].cells
    headers = ['日期', '行程安排', '住宿', '备注']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        
        # 设置表头格式
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
                run.font.bold = True
    
    # 示例数据行（第二行）- 这行将被Mail Merge变量替换
    sample_row = table.rows[1]
    sample_cells = sample_row.cells
    
    # 在第二行添加Mail Merge循环标记
    # 删除示例行，用Mail Merge循环替代
    sample_cells[0].text = '{% for day in days %}{{day.date}}'
    sample_cells[1].text = '{{day.activities_text}}'
    sample_cells[2].text = '{{day.hotel_name}}\n{{day.hotel_address}}\n{{day.hotel_phone}}'
    sample_cells[3].text = '{{day.highlights}}{% endfor %}'
    
    # 设置数据行格式
    for i, cell in enumerate(sample_cells):
        for paragraph in cell.paragraphs:
            if i == 0:  # 日期列居中
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:  # 其他列左对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(10)
    
    # 添加页脚信息
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run('生成时间：{{generated_date}} {{generated_time}}')
    footer_run.font.name = '微软雅黑'
    footer_run.font.size = Pt(9)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存模板
    template_path = "template_mailmerge_optimized.docx"
    doc.save(template_path)
    
    print(f"✅ Mail Merge模板创建成功: {template_path}")
    return template_path


def create_simple_mailmerge_template():
    """创建简化版Mail Merge模板 - 完全保持原模板格式"""
    
    print("📝 创建简化版Mail Merge模板...")
    print("   完全保持原模板的边距、字体、样式等设置...")
    
    # 从原模板复制
    if not os.path.exists("template.docx"):
        print("❌ 原模板文件不存在")
        return None
    
    # 读取原模板 - 这样可以保持所有原始设置（边距、字体、样式等）
    doc = Document("template.docx")
    
    # 查找并修改表格
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # 设置表格居中对齐
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 清除数据行，保留前2行（标题和表头）
        while len(table.rows) > 2:
            table._element.remove(table.rows[-1]._element)
        
        # 添加Mail Merge模板行
        new_row = table.add_row()
        cells = new_row.cells
        
        # 填入Mail Merge变量
        if len(cells) >= 4:
            cells[0].text = '{% for day in days %}{{day.date}}'
            cells[1].text = '{{day.activities_text}}'
            cells[2].text = '{{day.hotel_name}}\n{{day.hotel_address}}\n电话: {{day.hotel_phone}}'
            cells[3].text = '{{day.highlights}}{% endfor %}'
            
            # 设置格式
            for i, cell in enumerate(cells):
                for paragraph in cell.paragraphs:
                    if i == 0:  # 日期列居中
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:  # 其他列左对齐
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
    
    # 保存优化模板
    optimized_path = "template_mailmerge_simple.docx"
    doc.save(optimized_path)
    
    print(f"✅ 简化版Mail Merge模板创建成功: {optimized_path}")
    return optimized_path


def test_both_templates():
    """测试两个版本的模板"""
    print("=" * 60)
    print("创建Mail Merge模板")
    print("=" * 60)
    
    # 创建两个版本的模板
    template1 = create_mailmerge_template()
    template2 = create_simple_mailmerge_template()
    
    print()
    print("📄 已创建以下模板文件:")
    if template1:
        print(f"   1. {template1} - 全新创建的模板")
    if template2:
        print(f"   2. {template2} - 基于原模板优化的版本")
    
    print()
    print("🎯 推荐使用方式:")
    print("   - template_mailmerge_simple.docx 保持了原模板的所有格式")
    print("   - 可以直接用于Mail Merge生成，格式完美保持")
    
    return template2  # 返回推荐的模板


if __name__ == "__main__":
    test_both_templates()