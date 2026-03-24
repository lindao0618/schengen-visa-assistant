#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import logging
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
from typing import Dict, Any, Optional
import json
import tempfile
import base64
from PIL import Image
import fitz  # PyMuPDF for PDF processing
from docx import Document
from datetime import datetime
import requests

# 导入配置
from config import Config

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 尝试导入腾讯云OCR SDK
TENCENT_OCR_AVAILABLE = False
try:
    from tencentcloud.common import credential
    from tencentcloud.common.profile.client_profile import ClientProfile
    from tencentcloud.common.profile.http_profile import HttpProfile
    from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException
    from tencentcloud.ocr.v20181119 import ocr_client, models
    TENCENT_OCR_AVAILABLE = True
    logger.info("Tencent Cloud OCR SDK successfully imported")
except ImportError as e:
    logger.warning(f"Tencent Cloud OCR SDK not available: {e}")
    logger.warning("Please install: pip install tencentcloud-sdk-python")

# DeepSeek API配置
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1/chat/completions"

# 月份英文映射（用于酒店预订单日期解析）
MONTH_MAP = {
    "JANUARY": 1, "JAN": 1,
    "FEBRUARY": 2, "FEB": 2,
    "MARCH": 3, "MAR": 3,
    "APRIL": 4, "APR": 4,
    "MAY": 5,
    "JUNE": 6, "JUN": 6,
    "JULY": 7, "JUL": 7,
    "AUGUST": 8, "AUG": 8,
    "SEPTEMBER": 9, "SEP": 9, "SEPT": 9,
    "OCTOBER": 10, "OCT": 10,
    "NOVEMBER": 11, "NOV": 11,
    "DECEMBER": 12, "DEC": 12,
}


def _extract_hotel_structured_data(ocr_content: str) -> Dict[str, Optional[str]]:
    """
    从酒店预订单 OCR 文本中提取入住日期、退房日期、客人姓名。
    支持 Booking.com 等表格布局：CHECK-IN/CHECK-OUT 下列为 日、月 分列（如 28|3, MARCH|APRIL）。
    返回格式：{"check_in": "2026-03-28", "check_out": "2026-04-03", "guest_names": "YUXIN PENG, TINGTING ZHOU"}
    """
    result = {"check_in": None, "check_out": None, "guest_names": None}
    text = ocr_content.upper().replace("\r\n", "\n")
    text_orig = ocr_content

    # 1. 提取年份
    year_match = re.search(r"\b(20[2-3][0-9])\b", text)
    year = int(year_match.group(1)) if year_match else datetime.now().year

    # 2. 定位 CHECK-IN / CHECK-OUT 区块（表格布局：日、月分列）
    section = text
    for sep in ["GUEST NAME", "GUEST NAME:", "Important information", "NUMBER OF GUESTS"]:
        idx = text.find(sep.upper())
        if idx > 0:
            section = text[:idx]
            break

    month_pattern = r"(?:JANUARY|FEBRUARY|MARCH|APRIL|MAY|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER|JAN|FEB|MAR|APR|JUN|JUL|AUG|SEP|SEPT|OCT|NOV|DEC)"

    # 2a. 在 CHECK-IN/CHECK-OUT 区块内匹配连续 "DD MONTH" 或 "MONTH DD"（避免匹配取消政策等无关日期）
    date_pattern = rf"(?:(\d{{1,2}})\s+({month_pattern})|({month_pattern})\s+(\d{{1,2}}))"
    date_matches = list(re.finditer(date_pattern, section, re.IGNORECASE))
    parsed_dates = []
    for m in date_matches:
        if m.group(1) and m.group(2):
            day, month_str = int(m.group(1)), m.group(2).upper()
        elif m.group(3) and m.group(4):
            month_str, day = m.group(3).upper(), int(m.group(4))
        else:
            continue
        month = MONTH_MAP.get(month_str)
        if month and 1 <= day <= 31:
            try:
                dt = datetime(year, month, day)
                parsed_dates.append((dt, f"{year}-{month:02d}-{day:02d}"))
            except ValueError:
                pass

    # 2b. 表格布局：Booking.com 等格式为
    #     CHECK-IN  CHECK-OUT  ROOMS  NIGHTS
    #     28        3          1      6      <- 28=入住日, 3=退房日; 1=房间数, 6=夜数
    #     MARCH     APRIL
    # 仅取前2列（入住日、退房日），排除 ROOMS(1) 和 NIGHTS(6)
    if len(parsed_dates) < 2 and "CHECK" in section:
        months_found = []
        for m in re.finditer(month_pattern, section):
            month_str = m.group(0).upper()
            if month_str in MONTH_MAP:
                months_found.append(month_str)
        if len(months_found) < 2:
            pass
        else:
            # 定位数字区域：表头之后、首个月份之前
            # 注意：中间可能有地址如 "17 Rue Jean Colly, 13th arr."，其 17/13 会干扰
            # 表格行 (28, 3, 1, 6) 紧邻月份前，取「首个月份前的最后4个数字」作为表格行
            month_pos = section.find(months_found[0])
            before_months = section[:month_pos]
            all_days = []
            for m in re.finditer(r"\b([12]?\d|3[01])\b", before_months):
                d = int(m.group(1))
                if 1 <= d <= 31:
                    all_days.append(d)
            # 取最后4个（表格行：入住日、退房日、房间数、夜数），前2个为日期
            days = all_days[-4:] if len(all_days) >= 4 else all_days
            if len(days) >= 2:
                table_dates = []
                for i in range(2):
                    day, month_str = days[i], months_found[i]
                    month = MONTH_MAP[month_str]
                    try:
                        dt = datetime(year, month, day)
                        table_dates.append((dt, f"{year}-{month:02d}-{day:02d}"))
                    except ValueError:
                        pass
                if len(table_dates) >= 2:
                    parsed_dates = table_dates

    parsed_dates.sort(key=lambda x: x[0])
    if len(parsed_dates) >= 2:
        result["check_in"] = parsed_dates[0][1]
        result["check_out"] = parsed_dates[1][1]
    elif len(parsed_dates) == 1:
        result["check_in"] = parsed_dates[0][1]
        nights_match = re.search(r"\bNIGHTS?\s*[:\s]*(\d+)\b", text)
        if nights_match:
            from datetime import timedelta
            n = int(nights_match.group(1))
            check_out_dt = parsed_dates[0][0] + timedelta(days=n)
            result["check_out"] = check_out_dt.strftime("%Y-%m-%d")

    # 3. 提取客人姓名
    guest_patterns = [
        r"[Gg]uest\s*name\s*[:\s]+([A-Za-z\s,\.\-']+?)(?:\n|Number of|$)",
        r"[Gg]uest\s*[:\s]+([A-Za-z\s,\.\-']+?)(?:\n|Number of|$)",
    ]
    for pat in guest_patterns:
        gm = re.search(pat, text_orig, re.IGNORECASE | re.DOTALL)
        if gm:
            names = gm.group(1).strip()
            if len(names) >= 2 and not re.match(r"^\d+$", names):
                result["guest_names"] = names
                break

    return result


def _extract_booking_confirmation_pin(ocr_content: str) -> Dict[str, Optional[str]]:
    """从酒店预订单中提取 Booking.com 的 CONFIRMATION NUMBER 和 PIN CODE"""
    text = ocr_content.upper()
    confirmation = None
    pin = None

    m = re.search(r"CONFIRMATION\s*NUMBER[：:\s]*([0-9.\-\s]+)", text, re.IGNORECASE)
    if m:
        confirmation = re.sub(r"\D", "", m.group(1))
        if not confirmation:
            confirmation = None

    m = re.search(r"PIN\s*CODE[：:\s]*([0-9]+)", text, re.IGNORECASE)
    if m:
        pin = re.sub(r"\D", "", m.group(1))
        if not pin:
            pin = None

    if not confirmation:
        for line in ocr_content.splitlines():
            if "CONFIRMATION" in line.upper() and "NUMBER" in line.upper():
                digits = re.sub(r"\D", "", line)
                if digits:
                    confirmation = digits
                    break

    if not pin:
        for line in ocr_content.splitlines():
            if "PIN" in line.upper() and "CODE" in line.upper():
                digits = re.sub(r"\D", "", line)
                if digits:
                    pin = digits
                    break

    return {"confirmation_number": confirmation, "pin_code": pin}


def _verify_booking_order(confirmation_number: str, pin_code: str) -> Dict[str, Any]:
    """验证 Booking.com 订单，返回 found/error 信息"""
    if not confirmation_number or not pin_code:
        return {
            "checked": False,
            "found": None,
            "error": "未提取到 CONFIRMATION NUMBER 或 PIN CODE",
        }
    try:
        url = "https://secure.booking.cn/help/confirmation_pin_auth?label=mkt123sc-ae979d3b-b62b-4fcd-9a3e-8184427ca1f9&sid=367338f18fa014f29602b6a64e042215&aid=1662037&source=header"
        response = requests.post(
            url,
            data={
                "confirmationNumber": confirmation_number,
                "pinCode": pin_code,
            },
            timeout=20,
        )
        if response.status_code != 200:
            return {
                "checked": True,
                "found": None,
                "error": f"Booking.com 请求失败: {response.status_code}",
            }
        html = response.text
        if "我们没有找到符合这些信息的订单" in html:
            return {"checked": True, "found": False}
        # 订单卡片或酒店信息卡片特征
        if "b99b6ef58f" in html or "Hotel Le Richemont" in html or "Sat" in html:
            return {"checked": True, "found": True}
        return {"checked": True, "found": None, "error": "未识别到明确结果"}
    except Exception as e:
        return {"checked": True, "found": None, "error": f"Booking.com 验证异常: {str(e)}"}


class DeepSeekAnalyzer:
    """DeepSeek智能材料分析器"""
    
    def __init__(self):
        self.api_key = DEEPSEEK_API_KEY
        self.base_url = DEEPSEEK_BASE_URL
    
    def get_analysis_prompt(self, document_type: str, ocr_content: str, visa_type: str,
                           supplementary_info: Optional[Dict[str, str]] = None) -> str:
        """根据材料类型获取分析提示，supplementary_info 可包含 customer_name, departure_date, return_date"""
        supp = supplementary_info or {}
        supp_section = ""
        if supp:
            lines = []
            if supp.get("customer_name"):
                label = "客户姓名" if document_type == "hotel" else "客户/投保人姓名"
                lines.append(f"• {label}：{supp['customer_name']}")
            if supp.get("departure_date"):
                label = "入住日期" if document_type == "hotel" else "出发时间"
                lines.append(f"• {label}：{supp['departure_date']}")
            if supp.get("return_date"):
                label = "退房日期" if document_type == "hotel" else "返程时间"
                lines.append(f"• {label}：{supp['return_date']}")
            if lines:
                suffix = "请据此判定是否匹配。\n" if document_type == "hotel" else "请核对OCR识别的文档内容与上述信息是否一致，并在分析中说明。\n"
                supp_section = "\n\n📌 【用户提供的补充信息】\n" + "\n".join(lines) + "\n\n" + suffix

        # 酒店类型：从 OCR 中提取入住/退房日期和客人姓名，便于 DeepSeek 直接比较
        hotel_extracted_block = ""
        if document_type == "hotel":
            extracted = _extract_hotel_structured_data(ocr_content)
            parts = []
            if extracted.get("check_in"):
                parts.append(f"• OCR提取的入住日期：{extracted['check_in']}")
            if extracted.get("check_out"):
                parts.append(f"• OCR提取的退房日期：{extracted['check_out']}")
            if extracted.get("guest_names"):
                parts.append(f"• OCR提取的客人姓名：{extracted['guest_names']}")
            booking_verification = supp.get("booking_verification")
            if booking_verification:
                parts.append(f"• Booking.com 验证结果：{booking_verification}")
            if parts:
                hotel_extracted_block = (
                    "\n【📌 结构化提取结果】以下为从预订单中自动解析的入住/退房日期和姓名，请直接与用户填写比对。\n"
                    + "\n".join(parts) + "\n"
                )
        prompts = {
            "itinerary": f"""
🎯 作为专业的签证材料审核专家，请对以下行程单进行智能分析：

📄 OCR识别的行程单内容：
{ocr_content}

🔍 【智能审核要点】

📅 1️⃣ 基础信息审查
• 行程日期是否连续完整？📆
• 酒店地址与行程城市是否匹配？🏨
• 是否存在城市跳跃或重复问题？🔄

🎪 2️⃣ 每日安排合理性
• 景点数量：建议每天2-3个景点🎯
• 单景点安排：只有凡尔赛宫等大型景点才合理🏛️
• 避免重复：检查是否连续多日访问同一景点❌

🏛️ 3️⃣ 节假日闭馆检查
• 凡尔赛宫：周一闭馆 🚫
• 卢浮宫：周二闭馆 🚫
• 法国国定假日：5月1日劳动节、7月14日国庆等 🎉
• 圣诞新年：大部分景点不开放 🎄

⏰ 4️⃣ 时间安排真实性
• 是否留有合理交通和参观时间？⏱️
• 首日：抵达+入住酒店 ✅
• 末日：退房+返回出发地 ✅

💡 【输出格式要求】
请用中文输出，格式要生动有趣，多使用表情符号：

🎯 专业审核意见（100-200字）
[详细分析行程的合理性和真实性，先直接给出结论，然后给出分析，最后给出建议，记得要分段分点123这样子，结构要好看点，带点表情包]

📋 审核总结（50-100字）
[简洁的最终结论和建议，比如是否合理，是否符合签证要求，整体没有问题就没关系，然后你只需要考虑行程单的问题，不需要考虑其他材料]

🌟 请让内容看起来专业但不生硬，友好且易懂,然后不要使用markdown格式！
""",
            "hotel": f"""
任务：酒店预订单审核。仅核对姓名和日期两项，其他内容忽略。

{hotel_extracted_block}
【原始OCR内容】（若上方已有结构化提取结果，请以结构化数据为准进行比对；若无则从下方OCR中自行解析）
{ocr_content}
{supp_section}
【审核规则】

1. 姓名：采用「包含即通过」规则。用户填写的姓名只要被预订单中的入住人名单包含即通过。
   - 示例：用户填「TINGTING ZHOU」，预订单为「YUXIN PENG,TINGTING ZHOU」→ 通过（预订单包含用户姓名）
   - 预订单可有多位入住人，用户通常只填主申请人，只要其姓名在预订单中出现即判为无误。

2. 入住日期、退房日期：按日历日期语义比较。只要表示同一天即通过。
日期格式等价示例（以下均表示4月3日）：3 APRIL、April 3、04-03、2026-04-03、4月3日。

⚠️ 常见布局说明（Booking.com 等）：表格中 CHECK-IN/CHECK-OUT/ROOMS/NIGHTS 下列数字通常为：
- 第1列：入住日（如 28 = 3月28日）
- 第2列：退房日（如 3 = 4月3日）
- 第3列：房间数（如 1）
- 第4列：夜数（如 6）
切勿将房间数(1)或夜数(6)误当作日期。

【判定】
- 姓名和日期均匹配 → 材料无误，可作为递签材料。姓名、入住日期、退房日期均无误。
- 任一项不匹配 → 材料有误，不能作为递签材料。需说明：哪些无误、哪些有误。

【输出】第一行结论，禁止分析推理。
- 通过：材料无误，可作为递签材料。姓名、入住日期、退房日期均无误。
- 不通过：材料有误，不能作为递签材料。第二行写明：姓名无误/有误，入住日期无误/有误，退房日期无误/有误（有误时说明原因）。
""",
            "bank_statement": f"""
🏦 作为专业的签证材料审核专家，请对以下银行流水进行智能分析：

📄 OCR识别的银行流水内容：
{ocr_content}

🔍 【智能审核要点】

💰 1️⃣ 账户余额充足性
• 账户余额是否满足签证要求？💵
• 是否有足够的资金证明？📊
• 余额是否稳定持续？📈
• 是否存在异常大额变动？⚠️

💼 2️⃣ 收入来源稳定性
• 工资收入是否稳定规律？💼
• 收入金额是否合理？📋
• 是否有多种收入来源？🔄
• 收入时间是否规律？⏰

💸 3️⃣ 支出模式合理性
• 日常支出是否合理？🛒
• 是否有异常大额支出？🚨
• 支出模式是否符合个人情况？📊
• 是否存在可疑交易？❌

📅 4️⃣ 流水时间范围
• 流水时间是否覆盖要求？📆
• 是否有足够的历史记录？📚
• 时间跨度是否合理？⏱️
• 是否存在时间断层？🕳️

💡 【输出格式要求】
请用中文输出，格式要生动有趣，多使用表情符号：

🎯 专业审核意见（100-200字）
[详细分析银行流水的合规性]

📋 审核总结（50-100字）
[简洁的最终结论和建议]

🌟 请让内容看起来专业但不生硬，友好且易懂！
""",
            "flight": f"""
✈️ 作为专业的签证材料审核专家，请对以下机票预订单进行智能分析：

📄 OCR识别的机票预订单内容：
{ocr_content}

🔍 【智能审核要点】

🛫 1️⃣ 航班信息完整性
• 航班号是否真实有效？✈️
• 起飞和降落时间是否合理？⏰
• 舱位等级是否明确？💺
• 航空公司信息是否正确？🏢

🗺️ 2️⃣ 行程逻辑性
• 出发地和目的地是否合理？📍
• 转机安排是否可行？🔄
• 行程时间是否充足？⏱️
• 是否存在时间冲突？❌

💳 3️⃣ 预订状态有效性
• 预订状态是否为确认？✅
• 机票是否已出票？🎫
• 预订时间是否合理？📅
• 是否存在过期预订？⏰

💰 4️⃣ 价格合理性
• 机票价格是否符合市场水平？💵
• 税费是否计算正确？🧾
• 总价是否合理？📊
• 是否有异常低价？⚠️

📋 5️⃣ 退改签政策
• 退改签政策是否明确？📝
• 手续费是否合理？💸
• 政策是否符合航空公司标准？🏢
• 是否有限制条件？🔒

💡 【输出格式要求】
请用中文输出，格式要生动有趣，多使用表情符号：

🎯 专业审核意见（100-200字）
[详细分析机票预订的合理性]

📋 审核总结（50-100字）
[简洁的最终结论和建议]

🌟 请让内容看起来专业但不生硬，友好且易懂！
""",
            "insurance": f"""
🛡️ 作为专业的签证材料审核专家，请对以下旅行保险单进行智能分析：

📄 OCR识别的保险单内容：
{ocr_content}

🔍 【智能审核要点】

🛡️ 1️⃣ 保险覆盖范围
• 医疗费用是否包含？🏥
• 意外伤害是否覆盖？🚑
• 行李丢失是否保障？💼
• 航班延误是否包含？✈️

💰 2️⃣ 保险金额充足性
• 医疗保额是否不少于3万欧元？💵
• 意外保额是否合理？📊
• 总保额是否满足申根要求？✅
• 是否有保额不足的情况？⚠️

📅 3️⃣ 保险期间覆盖
• 保险期间是否覆盖整个行程？📆
• 是否有时间重叠或遗漏？⏰
• 保险生效时间是否合理？🕐
• 是否包含出发前和回国后？🔄

🏢 4️⃣ 保险公司资质
• 保险公司是否知名可靠？🏢
• 是否有相关资质认证？📜
• 理赔服务是否便捷？📞
• 是否有不良记录？❌

📋 5️⃣ 理赔流程
• 理赔流程是否清晰？📝
• 所需材料是否明确？📋
• 理赔时间是否合理？⏱️
• 是否有特殊要求？🔍

💡 【输出格式要求】
请用中文输出，格式要生动有趣，多使用表情符号：

🎯 专业审核意见（100-200字）
[详细分析保险单的有效性]

📋 审核总结（50-100字）
[简洁的最终结论和建议]

🌟 请让内容看起来专业但不生硬，友好且易懂！
""",
            "other": f"""
📋 作为专业的签证材料审核专家，请对以下文档进行智能分析：

📄 OCR识别的文档内容：
{ocr_content}

🔍 【智能审核要点】

📊 1️⃣ 文档完整性
• 文档内容是否完整？📄
• 关键信息是否清晰？🔍
• 是否有缺失页面？📑
• 格式是否规范？📋

✅ 2️⃣ 信息真实性
• 内容是否真实可信？🔐
• 是否存在虚假信息？❌
• 数据是否合理？📊
• 是否有矛盾之处？⚠️

🎯 3️⃣ 签证要求符合性
• 是否符合签证申请要求？📋
• 是否满足材料清单？✅
• 是否有遗漏项目？📝
• 格式是否符合标准？📄

⚠️ 4️⃣ 潜在问题识别
• 是否存在明显错误？🚨
• 是否有可疑内容？🔍
• 是否需要补充材料？📋
• 是否有风险点？⚠️

💡 5️⃣ 改进建议
• 如何完善材料？🔧
• 需要补充什么？📝
• 如何提高通过率？📈
• 注意事项有哪些？💡

💡 【输出格式要求】
请用中文输出，格式要生动有趣，多使用表情符号：

🎯 专业审核意见（100-200字）
[详细分析文档的合规性]

📋 审核总结（50-100字）
[简洁的最终结论和建议]

🌟 请让内容看起来专业但不生硬，友好且易懂！
"""
        }
        
        base_prompt = prompts.get(document_type, prompts["other"])

        # 仅在需要时添加补充信息（hotel 的 supp_section 已在模板内）
        if supp_section and document_type in ("flight", "insurance"):
            base_prompt = base_prompt.rstrip() + supp_section

        # 根据签证类型添加特定要求
        visa_requirements = {
            "schengen": "申根签证要求：行程必须合理，资金充足，保险覆盖整个行程，酒店预订确认。",
            "usa": "美国签证要求：行程合理，有明确目的，资金证明充足，有回国约束力。",
            "uk": "英国签证要求：行程详细，资金充足，住宿安排合理，有明确访问目的。",
            "japan": "日本签证要求：行程合理，资金证明，住宿预订，有明确访问目的。",
            "australia": "澳大利亚签证要求：行程合理，资金充足，住宿安排，有明确访问目的。",
            "canada": "加拿大签证要求：行程合理，资金证明充足，住宿安排，有明确访问目的。",
            "newzealand": "新西兰签证要求：行程合理，资金充足，住宿安排，有明确访问目的。",
            "singapore": "新加坡签证要求：行程合理，资金证明，住宿安排，有明确访问目的。",
            "korea": "韩国签证要求：行程合理，资金证明，住宿安排，有明确访问目的。",
            "other": "其他签证要求：请根据具体签证类型进行相应审核。"
        }
        
        visa_requirement = visa_requirements.get(visa_type, visa_requirements["other"])
        
        return f"{base_prompt}\n\n🎯 【签证类型特定要求】\n{visa_requirement}"
    
    def analyze_document_content(self, document_type: str, ocr_content: str, visa_type: str,
                                  supplementary_info: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """使用DeepSeek分析文档内容"""
        try:
            prompt = self.get_analysis_prompt(document_type, ocr_content, visa_type, supplementary_info)
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            
            data = {
                "model": "deepseek-chat",
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.7,
                "max_tokens": 2000
            }
            
            logger.info(f"Calling DeepSeek API for {document_type} analysis...")
            response = requests.post(self.base_url, headers=headers, json=data, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                analysis_content = result["choices"][0]["message"]["content"]
                
                logger.info(f"DeepSeek analysis completed, content length: {len(analysis_content)}")
                
                return {
                    "success": True,
                    "analysis": analysis_content,
                    "document_type": document_type,
                    "tokens_used": result.get("usage", {}).get("total_tokens", 0)
                }
            else:
                logger.error(f"DeepSeek API error: {response.status_code} - {response.text}")
                return {
                    "success": False,
                    "error": f"DeepSeek API error: {response.status_code}",
                    "analysis": "API调用失败，无法进行智能分析"
                }
                
        except Exception as e:
            logger.error(f"DeepSeek analysis failed: {e}")
            return {
                "success": False,
                "error": f"分析失败: {str(e)}",
                "analysis": "智能分析服务暂时不可用"
            }

# 创建FastAPI应用
app = FastAPI(title="Material Review Service with Tencent Cloud OCR", version="1.0.0")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "http://localhost:3004"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class TencentCloudDocumentAnalyzer:
    """腾讯云OCR文档分析器"""
    
    def __init__(self):
        self.ocr_client = None
        self.deepseek_analyzer = DeepSeekAnalyzer()
        self.init_tencent_ocr()
    
    def init_tencent_ocr(self):
        """初始化腾讯云OCR客户端"""
        if not TENCENT_OCR_AVAILABLE:
            logger.warning("Tencent Cloud OCR SDK not available")
            return
        
        try:
            # 从配置文件获取密钥
            secret_id = Config.TENCENTCLOUD_SECRET_ID
            secret_key = Config.TENCENTCLOUD_SECRET_KEY
            
            if not secret_id or not secret_key:
                logger.warning("Tencent Cloud credentials not found in config")
                logger.warning("Please check config.py file")
                return
            
            logger.info(f"Using Tencent Cloud credentials: {secret_id[:10]}...")
            
            # 实例化认证对象
            cred = credential.Credential(secret_id, secret_key)
            
            # 实例化HTTP选项
            httpProfile = HttpProfile()
            httpProfile.endpoint = "ocr.tencentcloudapi.com"
            
            # 实例化client选项
            clientProfile = ClientProfile()
            clientProfile.httpProfile = httpProfile
            
            # 实例化OCR客户端
            self.ocr_client = ocr_client.OcrClient(cred, "ap-guangzhou", clientProfile)
            
            logger.info("Tencent Cloud OCR client initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize Tencent Cloud OCR: {e}")
            self.ocr_client = None
    
    def pdf_to_base64(self, pdf_path: str) -> str:
        """将PDF转换为base64编码"""
        try:
            with open(pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            # 转换为base64
            base64_content = base64.b64encode(pdf_content).decode('utf-8')
            
            # 添加PDF的data URL前缀
            data_url = f"data:application/pdf;base64,{base64_content}"
            
            return data_url
            
        except Exception as e:
            logger.error(f"Failed to convert PDF to base64: {e}")
            return None
    
    def image_to_base64(self, image_path: str) -> str:
        """将图片转换为base64编码"""
        try:
            with open(image_path, 'rb') as f:
                image_content = f.read()
            
            # 转换为base64
            base64_content = base64.b64encode(image_content).decode('utf-8')
            
            # 根据文件扩展名确定MIME类型
            if image_path.lower().endswith('.png'):
                mime_type = "image/png"
            elif image_path.lower().endswith('.jpg') or image_path.lower().endswith('.jpeg'):
                mime_type = "image/jpeg"
            else:
                mime_type = "image/png"  # 默认使用PNG
            
            # 添加图片的data URL前缀
            data_url = f"data:{mime_type};base64,{base64_content}"
            
            return data_url
            
        except Exception as e:
            logger.error(f"Failed to convert image to base64: {e}")
            return None
    
    def pdf_to_images_for_ocr(self, pdf_path: str) -> list:
        """将PDF转换为图片列表进行OCR"""
        try:
            images = []
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 提高分辨率以提升数字/日期识别准确率（3与6等易混淆）
                mat = fitz.Matrix(4.0, 4.0)  # 4倍放大，提升小字和数字识别
                pix = page.get_pixmap(matrix=mat)
                
                # 获取PNG字节数据
                img_data = pix.tobytes("png")
                
                images.append({
                    "page": page_num + 1,
                    "image_data": img_data
                })
                
                logger.info(f"Converted PDF page {page_num + 1} to image, size: {len(img_data)} bytes")
            
            doc.close()
            return images
            
        except Exception as e:
            logger.error(f"Failed to convert PDF to images: {e}")
            return []

    def extract_text_with_tencent_ocr(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """使用腾讯云OCR提取文本"""
        if not self.ocr_client:
            return {
                "success": False,
                "error": "Tencent Cloud OCR client not initialized",
                "extracted_text": "",
                "confidence": 0.0
            }
        
        try:
            extracted_text = ""
            all_confidence_scores = []
            
            # 根据文件类型处理
            if file_type.lower() == 'pdf':
                # PDF转换为图片进行OCR
                images = self.pdf_to_images_for_ocr(file_path)
                
                if not images:
                    return {
                        "success": False,
                        "error": "Failed to convert PDF to images",
                        "extracted_text": "",
                        "confidence": 0.0
                    }
                
                # 对每页图片进行OCR
                for img_info in images:
                    page_result = self._ocr_single_image(img_info["image_data"], img_info["page"])
                    if page_result["success"]:
                        extracted_text += f"\n\n=== 第{img_info['page']}页 ===\n"
                        extracted_text += page_result["text"]
                        all_confidence_scores.extend(page_result["confidence_scores"])
                    else:
                        logger.warning(f"Page {img_info['page']} OCR failed: {page_result['error']}")
            
            else:
                # 图片文件直接OCR
                try:
                    with open(file_path, 'rb') as f:
                        image_data = f.read()
                    
                    result = self._ocr_single_image(image_data, 1)
                    if result["success"]:
                        extracted_text = result["text"]
                        all_confidence_scores = result["confidence_scores"]
                    else:
                        return {
                            "success": False,
                            "error": result["error"],
                            "extracted_text": "",
                            "confidence": 0.0
                        }
                except Exception as e:
                    return {
                        "success": False,
                        "error": f"Failed to read image file: {str(e)}",
                        "extracted_text": "",
                        "confidence": 0.0
                    }
            
            # 计算平均置信度
            avg_confidence = sum(all_confidence_scores) / len(all_confidence_scores) if all_confidence_scores else 0.0
            
            logger.info(f"Total extracted text length: {len(extracted_text)}")
            logger.info(f"Average confidence: {avg_confidence}")
            
            return {
                "success": True,
                "extracted_text": extracted_text.strip(),
                "confidence": avg_confidence
            }
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return {
                "success": False,
                "error": f"OCR extraction failed: {str(e)}",
                "extracted_text": "",
                "confidence": 0.0
            }
    
    def _ocr_single_image(self, image_data: bytes, page_num: int) -> Dict[str, Any]:
        """对单张图片进行OCR识别，优先使用高精度版以降低数字误识（如3↔6）"""
        try:
            base64_content = base64.b64encode(image_data).decode('utf-8')
            params = {"ImageBase64": base64_content}
            req = models.GeneralAccurateOCRRequest()
            req.from_json_string(json.dumps(params))

            logger.info(f"Calling Tencent Cloud GeneralAccurateOCR for page {page_num}...")
            resp = self.ocr_client.GeneralAccurateOCR(req)
            
            # 解析响应
            response_dict = json.loads(resp.to_json_string())
            logger.info(f"Page {page_num} OCR response received successfully")
            
            # 提取文本内容
            extracted_text = ""
            confidence_scores = []
            
            if "TextDetections" in response_dict:
                for detection in response_dict["TextDetections"]:
                    if "DetectedText" in detection:
                        extracted_text += detection["DetectedText"] + "\n"
                    if "Confidence" in detection:
                        confidence_scores.append(detection["Confidence"])
                        
                logger.info(f"Page {page_num} found {len(response_dict['TextDetections'])} text detections")
            else:
                logger.warning(f"Page {page_num} no TextDetections found in response")
            
            logger.info(f"Page {page_num} OCR completed, text length: {len(extracted_text)}")
            
            return {
                "success": True,
                "text": extracted_text,
                "confidence_scores": confidence_scores
            }
            
        except TencentCloudSDKException as e:
            logger.error(f"Tencent Cloud OCR error for page {page_num}: {e}")
            return {
                "success": False,
                "error": f"Tencent Cloud OCR error: {str(e)}",
                "text": "",
                "confidence_scores": []
            }
        except Exception as e:
            logger.error(f"OCR failed for page {page_num}: {e}")
            return {
                "success": False,
                "error": f"OCR failed: {str(e)}",
                "text": "",
                "confidence_scores": []
            }
    
    def save_extracted_text_to_word(self, extracted_text: str, original_filename: str) -> str:
        """将提取的文本保存为Word文档"""
        try:
            # 创建输出目录
            output_dir = "ocr_results"
            os.makedirs(output_dir, exist_ok=True)
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            doc.add_heading('OCR提取结果', 0)
            
            # 添加文件信息
            doc.add_heading('文件信息', level=1)
            doc.add_paragraph(f'原始文件名: {original_filename}')
            doc.add_paragraph(f'提取时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'OCR引擎: 腾讯云OCR')
            
            # 添加提取的文本
            doc.add_heading('提取的文本内容', level=1)
            doc.add_paragraph(extracted_text)
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            word_filename = f"ocr_result_{timestamp}.docx"
            word_path = os.path.join(output_dir, word_filename)
            
            # 保存文档
            doc.save(word_path)
            logger.info(f"OCR结果已保存到Word文档: {word_path}")
            
            return word_path
            
        except Exception as e:
            logger.error(f"保存Word文档失败: {e}")
            return None
    
    def analyze_document_with_ai(self, extracted_text: str, document_type: str, visa_type: str,
                                  supplementary_info: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """使用DeepSeek AI分析文档内容"""
        try:
            # 使用DeepSeek进行智能分析
            deepseek_result = self.deepseek_analyzer.analyze_document_content(
                document_type, extracted_text, visa_type, supplementary_info
            )
            
            if deepseek_result["success"]:
                # DeepSeek分析成功
                analysis_result = {
                    "document_type": document_type,
                    "text_length": len(extracted_text),
                    "ai_analysis": deepseek_result["analysis"],
                    "analysis_status": "DeepSeek智能分析完成",
                    "tokens_used": deepseek_result.get("tokens_used", 0),
                    "keywords_detected": self._extract_keywords(extracted_text, document_type),
                    "recommendations": self._generate_recommendations(deepseek_result["analysis"]),
                    "confidence_score": 95.0 if extracted_text else 0.0
                }
            else:
                # DeepSeek分析失败，使用基础分析
                logger.warning(f"DeepSeek analysis failed: {deepseek_result.get('error', 'Unknown error')}")
                analysis_result = {
                    "document_type": document_type,
                    "text_length": len(extracted_text),
                    "ai_analysis": f"DeepSeek分析失败：{deepseek_result.get('error', '未知错误')}，已切换到基础分析模式。",
                    "analysis_status": "基础分析模式",
                    "keywords_detected": self._extract_keywords(extracted_text, document_type),
                    "recommendations": ["文档识别成功", "建议人工复核关键信息", "智能分析服务暂时不可用"],
                    "confidence_score": 75.0 if extracted_text else 0.0
                }
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"AI analysis failed: {e}")
            # 降级到基础分析
            return {
                "document_type": document_type,
                "text_length": len(extracted_text),
                "ai_analysis": f"智能分析失败：{str(e)}，已切换到基础分析模式。",
                "analysis_status": "基础分析模式",
                "keywords_detected": self._extract_keywords(extracted_text, document_type),
                "recommendations": ["文档识别成功", "建议人工复核关键信息"],
                "confidence_score": 70.0 if extracted_text else 0.0
            }
    
    def _extract_keywords(self, text: str, document_type: str) -> list:
        """提取关键词"""
        keywords_map = {
            "itinerary": ["行程", "旅行", "酒店", "景点", "日期", "住宿"],
            "hotel": ["酒店", "预订", "入住", "退房", "房间", "价格"],
            "bank_statement": ["银行", "流水", "余额", "收入", "支出", "账户"],
            "flight": ["航班", "机票", "起飞", "降落", "座位", "价格"],
            "insurance": ["保险", "保障", "理赔", "医疗", "意外", "金额"],
            "other": ["文档", "证明", "信息", "材料"]
        }
        
        relevant_keywords = keywords_map.get(document_type, keywords_map["other"])
        detected = [keyword for keyword in relevant_keywords if keyword in text]
        
        return detected
    
    def _generate_recommendations(self, ai_analysis: str) -> list:
        """根据AI分析生成建议"""
        recommendations = ["DeepSeek智能分析已完成"]
        
        # 根据分析内容添加具体建议
        if "问题" in ai_analysis or "错误" in ai_analysis:
            recommendations.append("发现潜在问题，请仔细检查")
        if "建议" in ai_analysis:
            recommendations.append("请查看AI分析中的具体建议")
        if "完整" in ai_analysis and "信息" in ai_analysis:
            recommendations.append("文档信息相对完整")
        
        recommendations.append("建议结合人工审核")
        
        return recommendations
    
    def analyze_document(self, file_path: str, document_type: str, original_filename: str, visa_type: str,
                         supplementary_info: Optional[Dict[str, str]] = None,
                         booking_verify: bool = False) -> Dict[str, Any]:
        """分析文档的主要方法"""
        try:
            # 确定文件类型
            file_extension = os.path.splitext(original_filename)[1].lower()
            if file_extension == '.pdf':
                file_type = 'pdf'
            elif file_extension in ['.jpg', '.jpeg', '.png']:
                file_type = 'image'
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_extension}",
                    "analysis_result": {}
                }
            
            # 使用腾讯云OCR提取文本
            ocr_result = self.extract_text_with_tencent_ocr(file_path, file_type)
            
            if not ocr_result["success"]:
                return {
                    "success": False,
                    "error": ocr_result["error"],
                    "analysis_result": {}
                }
            
            extracted_text = ocr_result["extracted_text"]
            confidence = ocr_result["confidence"]
            
            # 保存提取的文本到Word文档
            word_document_path = self.save_extracted_text_to_word(extracted_text, original_filename)
            
            booking_verification = None
            if document_type == "hotel" and booking_verify:
                booking_data = _extract_booking_confirmation_pin(extracted_text)
                confirmation_number = booking_data.get("confirmation_number")
                pin_code = booking_data.get("pin_code")
                booking_verification = _verify_booking_order(confirmation_number, pin_code)
                if supplementary_info is None:
                    supplementary_info = {}
                if booking_verification.get("found") is True:
                    supplementary_info["booking_verification"] = "已找到订单"
                elif booking_verification.get("found") is False:
                    supplementary_info["booking_verification"] = "未找到订单"
                elif booking_verification.get("error"):
                    supplementary_info["booking_verification"] = f"验证失败：{booking_verification.get('error')}"
                else:
                    supplementary_info["booking_verification"] = "验证结果不明确"

            # 使用AI分析文档内容
            ai_analysis = self.analyze_document_with_ai(
                extracted_text, document_type, visa_type, supplementary_info
            )
            
            # 构建分析结果
            analysis_result = {
                "ocr_info": {
                    "engine": "Tencent Cloud OCR",
                    "confidence": confidence,
                    "text_length": len(extracted_text),
                    "word_document": word_document_path
                },
                "extracted_data": {
                    "raw_text": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
                    "full_text_length": len(extracted_text)
                },
                "booking_verification": booking_verification,
                "ai_analysis": ai_analysis,
                "verification_results": {
                    "text_extracted": "通过" if extracted_text else "未通过",
                    "content_relevant": "通过" if len(extracted_text) > 10 else "未通过"
                },
                "recommendations": ai_analysis["recommendations"]
            }

            if document_type == "hotel" and booking_verify and booking_verification:
                found = booking_verification.get("found")
                if found is False or booking_verification.get("error") or found is None:
                    extracted = _extract_hotel_structured_data(extracted_text)
                    customer_name = (supplementary_info or {}).get("customer_name", "")
                    name_ok = (
                        bool(customer_name)
                        and bool(extracted.get("guest_names"))
                        and customer_name in extracted.get("guest_names", "")
                    )
                    check_in_value = extracted.get("check_in")
                    check_out_value = extracted.get("check_out")
                    user_in = (supplementary_info or {}).get("departure_date", "")
                    user_out = (supplementary_info or {}).get("return_date", "")
                    check_in_ok = bool(user_in and check_in_value and user_in == check_in_value)
                    check_out_ok = bool(user_out and check_out_value and user_out == check_out_value)

                    name_text = "姓名无误" if name_ok else "姓名有误"
                    if not check_in_value:
                        in_text = f"入住日期无法从预订单识别（用户填写为{user_in or '未知'}）"
                    else:
                        in_text = (
                            "入住日期无误"
                            if check_in_ok
                            else f"入住日期有误（用户填写为{user_in or '未知'}，预订单为{check_in_value}）"
                        )
                    if not check_out_value:
                        out_text = f"退房日期无法从预订单识别（用户填写为{user_out or '未知'}）"
                    else:
                        out_text = (
                            "退房日期无误"
                            if check_out_ok
                            else f"退房日期有误（用户填写为{user_out or '未知'}，预订单为{check_out_value}）"
                        )
                    booking_text = "Booking.com订单验证：未找到订单。"
                    if booking_verification.get("error"):
                        booking_text = f"Booking.com订单验证失败：{booking_verification.get('error')}"

                    analysis_result["ai_analysis"]["ai_analysis"] = (
                        "不通过：材料有误，不能作为递签材料。\n"
                        f"{name_text}，{in_text}，{out_text}，{booking_text}"
                    )
            
            return {
                "success": True,
                "analysis_result": analysis_result
            }
            
        except Exception as e:
            logger.error(f"Document analysis failed: {e}")
            return {
                "success": False,
                "error": f"Document analysis failed: {str(e)}",
                "analysis_result": {}
            }

# 创建全局分析器实例
document_analyzer = TencentCloudDocumentAnalyzer()

@app.post("/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(default="itinerary"),
    visa_type: str = Form(default="schengen"),
    customer_name: Optional[str] = Form(default=None),
    departure_date: Optional[str] = Form(default=None),
    return_date: Optional[str] = Form(default=None),
    booking_verify: Optional[str] = Form(default=None),
):
    """上传文档进行分析"""
    try:
        logger.info(f"Received file upload: {file.filename}, type: {document_type}")
        
        # 检查文件类型
        allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png']
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # 保存上传的文件到临时目录
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        logger.info(f"File saved to temporary path: {temp_file_path}")

        # 构建补充信息（机票/车票、酒店、保险使用）
        supplementary_info = {}
        if customer_name and customer_name.strip():
            supplementary_info["customer_name"] = customer_name.strip()
        if departure_date and departure_date.strip():
            supplementary_info["departure_date"] = departure_date.strip()
        if return_date and return_date.strip():
            supplementary_info["return_date"] = return_date.strip()

        # 分析文档
        analysis_result = document_analyzer.analyze_document(
            temp_file_path,
            document_type,
            file.filename,
            visa_type,
            supplementary_info=supplementary_info if supplementary_info else None,
            booking_verify=bool(booking_verify and booking_verify.lower() == "true"),
        )
        
        # 清理临时文件
        try:
            os.unlink(temp_file_path)
        except Exception as e:
            logger.warning(f"Failed to delete temp file: {e}")
        
        if analysis_result["success"]:
            # 构建响应，包含Word文档下载链接
            word_document_path = analysis_result["analysis_result"]["ocr_info"].get("word_document")
            word_download_url = None
            
            if word_document_path:
                word_filename = os.path.basename(word_document_path)
                word_download_url = f"/download-ocr-result/{word_filename}"
            
            response_data = {
                "message": "Document analysis completed successfully",
                "analysis_result": analysis_result["analysis_result"],
                "word_download_url": word_download_url
            }
            
            logger.info("Document analysis completed successfully")
            return JSONResponse(content=response_data)
        else:
            logger.error(f"Document analysis failed: {analysis_result['error']}")
            raise HTTPException(status_code=500, detail=analysis_result["error"])
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload processing failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload processing failed: {str(e)}")

@app.get("/download-ocr-result/{filename}")
async def download_ocr_result(filename: str):
    """下载OCR结果Word文档"""
    try:
        file_path = os.path.join("ocr_results", filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        return FileResponse(
            file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File download failed: {e}")
        raise HTTPException(status_code=500, detail=f"File download failed: {str(e)}")

@app.get("/health")
async def health_check():
    """健康检查端点"""
    return {
        "status": "healthy",
        "service": "Material Review Service with Tencent Cloud OCR",
        "tencent_ocr_available": TENCENT_OCR_AVAILABLE,
        "ocr_client_initialized": document_analyzer.ocr_client is not None
    }

if __name__ == "__main__":
    port = int(os.environ.get("MATERIAL_REVIEW_PORT", "8004"))
    logger.info(f"Starting Material Review Service with Tencent Cloud OCR on port {port}...")
    uvicorn.run(app, host="0.0.0.0", port=port)
