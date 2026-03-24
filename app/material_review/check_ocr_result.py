#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
查看OCR结果工具
"""

import os
from docx import Document
import json

def read_ocr_word_document(docx_path: str):
    """读取OCR结果Word文档内容"""
    try:
        doc = Document(docx_path)
        
        print(f"=== OCR文档内容: {os.path.basename(docx_path)} ===")
        print()
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"段落 {i+1}: {text}")
        
        print()
        print("=== 文档分析 ===")
        total_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        print(f"总字符数: {len(total_text)}")
        print(f"总段落数: {len([p for p in doc.paragraphs if p.text.strip()])}")
        
        if len(total_text) == 0:
            print("⚠️ 警告: 文档中没有提取到任何文本内容！")
        else:
            print("✅ 文档包含文本内容")
        
        return total_text
        
    except Exception as e:
        print(f"❌ 读取文档失败: {e}")
        return ""

def main():
    """主函数"""
    results_dir = "ocr_results"
    
    if not os.path.exists(results_dir):
        print(f"❌ 结果目录不存在: {results_dir}")
        return
    
    # 列出所有Word文档
    docx_files = [f for f in os.listdir(results_dir) if f.endswith('.docx') and not f.startswith('~')]
    
    if not docx_files:
        print("❌ 没有找到OCR结果文档")
        return
    
    # 按修改时间排序，最新的在前
    docx_files.sort(key=lambda x: os.path.getmtime(os.path.join(results_dir, x)), reverse=True)
    
    print("=== 可用的OCR结果文档 ===")
    for i, filename in enumerate(docx_files):
        filepath = os.path.join(results_dir, filename)
        mtime = os.path.getmtime(filepath)
        import datetime
        mtime_str = datetime.datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")
        print(f"{i+1}. {filename} (修改时间: {mtime_str})")
    
    print()
    
    # 读取最新的文档
    latest_file = os.path.join(results_dir, docx_files[0])
    print(f"正在读取最新文档: {docx_files[0]}")
    print("=" * 50)
    
    content = read_ocr_word_document(latest_file)
    
    return content

if __name__ == "__main__":
    main()