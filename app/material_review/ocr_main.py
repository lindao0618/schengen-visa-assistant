#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import logging
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
from typing import Dict, Any
import json
import tempfile
from PIL import Image
import fitz  # PyMuPDF for PDF processing
from docx import Document
from datetime import datetime

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 尝试导入PaddleOCR，使用最基础的功能
PADDLE_OCR_AVAILABLE = False
try:
    from paddleocr import PaddleOCR
    PADDLE_OCR_AVAILABLE = True
    logger.info("PaddleOCR successfully imported")
except ImportError as e:
    logger.warning(f"PaddleOCR not available: {e}")

# 创建FastAPI应用
app = FastAPI(title="Material Review Service with OCR", version="1.0.0")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3004"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class OptimizedDocumentAnalyzer:
    """优化的文档分析器，使用基础PaddleOCR功能"""
    
    def __init__(self):
        self.ocr_engine = None
        self.available = False
        
        if PADDLE_OCR_AVAILABLE:
            try:
                # 只使用最基础的OCR功能，最小化参数
                self.ocr_engine = PaddleOCR(
                    lang='ch'  # 只指定语言，其他使用默认值
                )
                self.available = True
                logger.info("Basic PaddleOCR initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize PaddleOCR: {e}")
                self.ocr_engine = None
                self.available = False
        else:
            logger.warning("PaddleOCR not available, using fallback mode")
    
    def pdf_to_images(self, pdf_path: str) -> list:
        """将PDF转换为图片列表，优化分辨率和质量"""
        images = []
        try:
            pdf_document = fitz.open(pdf_path)
            logger.info(f"PDF has {len(pdf_document)} pages")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # 大幅提高分辨率以获得更好的OCR效果
                mat = fitz.Matrix(4.0, 4.0)  # 4倍缩放，提高清晰度
                pix = page.get_pixmap(matrix=mat, alpha=False)  # 移除alpha通道
                img_data = pix.tobytes("png")
                
                # 保存临时图片
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                    tmp_file.write(img_data)
                    images.append(tmp_file.name)
                    logger.info(f"Created high-res image for page {page_num + 1}: {tmp_file.name}")
            
            pdf_document.close()
            return images
        except Exception as e:
            logger.error(f"Failed to convert PDF to images: {e}")
            return []
    
    def extract_text_with_ocr(self, file_path: str) -> Dict[str, Any]:
        """使用基础OCR提取文本"""
        if not self.available:
            return {"error": "OCR not available", "text": ""}
        
        try:
            # 检查文件类型
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                # PDF文件：转换为图片后OCR
                image_paths = self.pdf_to_images(file_path)
                if not image_paths:
                    return {"error": "Failed to convert PDF", "text": ""}
                
                all_text = []
                for page_idx, img_path in enumerate(image_paths):
                    try:
                        logger.info(f"Starting OCR for page {page_idx + 1}: {img_path}")
                        result = self.ocr_engine.predict(img_path)
                        
                        logger.info(f"OCR result type: {type(result)}, length: {len(result) if result else 0}")
                        
                        if result and len(result) > 0 and result[0]:
                            page_text = []
                            logger.info(f"Found {len(result[0])} text blocks on page {page_idx + 1}")
                            
                            for line_idx, line in enumerate(result[0]):
                                if len(line) >= 2:  # 确保有坐标和文本
                                    text_content = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                                    confidence = line[1][1] if isinstance(line[1], (list, tuple)) and len(line[1]) > 1 else 0.0
                                    
                                    logger.info(f"Line {line_idx}: '{text_content}' (confidence: {confidence})")
                                    
                                    if text_content and text_content.strip():
                                        page_text.append(text_content.strip())
                            
                            if page_text:
                                page_content = '\n'.join(page_text)
                                all_text.append(page_content)
                                logger.info(f"Page {page_idx + 1} extracted {len(page_text)} text lines")
                            else:
                                logger.warning(f"No valid text extracted from page {page_idx + 1}")
                        else:
                            logger.warning(f"No OCR results for page {page_idx + 1}")
                        
                        # 清理临时文件
                        try:
                            os.unlink(img_path)
                        except:
                            pass
                            
                    except Exception as e:
                        logger.error(f"OCR failed for page {page_idx + 1} ({img_path}): {e}")
                        continue
                
                return {
                    "text": '\n\n'.join(all_text),
                    "pages": len(image_paths),
                    "method": "PDF->Images->OCR"
                }
            
            else:
                # 图片文件：直接OCR
                logger.info(f"Starting direct OCR for image: {file_path}")
                result = self.ocr_engine.predict(file_path)
                
                logger.info(f"Direct OCR result type: {type(result)}, length: {len(result) if result else 0}")
                
                if result and len(result) > 0 and result[0]:
                    text_lines = []
                    logger.info(f"Found {len(result[0])} text blocks in image")
                    
                    for line_idx, line in enumerate(result[0]):
                        if len(line) >= 2:  # 确保有坐标和文本
                            text_content = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                            confidence = line[1][1] if isinstance(line[1], (list, tuple)) and len(line[1]) > 1 else 0.0
                            
                            logger.info(f"Line {line_idx}: '{text_content}' (confidence: {confidence})")
                            
                            if text_content and text_content.strip():
                                text_lines.append(text_content.strip())
                    
                    return {
                        "text": '\n'.join(text_lines),
                        "method": "Direct OCR"
                    }
                else:
                    logger.warning("No OCR results for direct image")
                    return {"text": "", "method": "Direct OCR", "note": "No text detected"}
        
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return {"error": str(e), "text": ""}
    
    def analyze_document_with_ai(self, extracted_text: str, analysis_type: str, file_info: dict) -> Dict[str, Any]:
        """基于提取的文本进行智能分析"""
        text_length = len(extracted_text)
        
        # 简单的关键词分析
        keywords_found = []
        
        if analysis_type == "itinerary":
            keywords = ["行程", "日期", "酒店", "航班", "火车", "景点", "巴黎", "罗马", "柏林", "伦敦"]
            destinations = []
            
            for keyword in keywords:
                if keyword in extracted_text:
                    keywords_found.append(keyword)
                    if keyword in ["巴黎", "罗马", "柏林", "伦敦"]:
                        destinations.append(keyword)
            
            analysis_result = {
                "extracted_data": {
                    "document_type": "行程单",
                    "file_name": file_info["name"],
                    "file_size": file_info["size"],
                    "text_length": text_length,
                    "keywords_found": keywords_found,
                    "destinations": destinations if destinations else ["未识别"],
                    "raw_text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
                },
                "verification_results": {
                    "text_extracted": "通过" if text_length > 0 else "失败",
                    "keywords_detected": "通过" if keywords_found else "未检测到",
                    "content_relevant": "通过" if any(k in extracted_text for k in ["行程", "日期", "酒店"]) else "需要检查"
                },
                "recommendations": [
                    f"成功提取{text_length}个字符的文本内容",
                    f"检测到{len(keywords_found)}个相关关键词",
                    "建议人工审核文本内容的准确性" if text_length > 0 else "建议检查文件质量或重新上传"
                ],
                "confidence_score": min(0.6 + (len(keywords_found) * 0.05), 0.95) if text_length > 0 else 0.1
            }
        
        elif analysis_type == "hotel_booking":
            keywords = ["酒店", "预订", "确认", "入住", "退房", "客人", "房间"]
            
            for keyword in keywords:
                if keyword in extracted_text:
                    keywords_found.append(keyword)
            
            analysis_result = {
                "extracted_data": {
                    "document_type": "酒店预订单", 
                    "file_name": file_info["name"],
                    "file_size": file_info["size"],
                    "text_length": text_length,
                    "keywords_found": keywords_found,
                    "raw_text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
                },
                "verification_results": {
                    "text_extracted": "通过" if text_length > 0 else "失败",
                    "booking_keywords": "通过" if any(k in extracted_text for k in ["酒店", "预订"]) else "需要检查",
                    "content_relevant": "通过" if keywords_found else "需要检查"
                },
                "recommendations": [
                    f"成功提取{text_length}个字符的文本内容",
                    f"检测到{len(keywords_found)}个酒店相关关键词",
                    "建议验证预订确认号和日期信息"
                ],
                "confidence_score": min(0.7 + (len(keywords_found) * 0.04), 0.95) if text_length > 0 else 0.1
            }
        
        else:  # bank_statement
            keywords = ["银行", "流水", "余额", "收入", "支出", "账户", "明细"]
            
            for keyword in keywords:
                if keyword in extracted_text:
                    keywords_found.append(keyword)
            
            analysis_result = {
                "extracted_data": {
                    "document_type": "银行流水",
                    "file_name": file_info["name"], 
                    "file_size": file_info["size"],
                    "text_length": text_length,
                    "keywords_found": keywords_found,
                    "raw_text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
                },
                "verification_results": {
                    "text_extracted": "通过" if text_length > 0 else "失败",
                    "bank_keywords": "通过" if any(k in extracted_text for k in ["银行", "流水", "余额"]) else "需要检查",
                    "content_relevant": "通过" if keywords_found else "需要检查"
                },
                "recommendations": [
                    f"成功提取{text_length}个字符的文本内容",
                    f"检测到{len(keywords_found)}个银行相关关键词", 
                    "建议检查余额和交易记录的完整性"
                ],
                "confidence_score": min(0.6 + (len(keywords_found) * 0.06), 0.95) if text_length > 0 else 0.1
            }
        
        return analysis_result
    
    def save_extracted_text_to_word(self, extracted_text: str, file_name: str) -> str:
        """将提取的文本保存为Word文档"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            doc.add_heading(f'OCR提取结果 - {file_name}', 0)
            
            # 添加基本信息
            doc.add_heading('文档信息', level=1)
            doc.add_paragraph(f'原文件名: {file_name}')
            doc.add_paragraph(f'提取时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'文本长度: {len(extracted_text)} 字符')
            
            # 添加提取的文本
            doc.add_heading('提取的文本内容', level=1)
            if extracted_text.strip():
                # 按段落分割文本
                paragraphs = extracted_text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph('未提取到文本内容')
            
            # 保存文档
            output_dir = "ocr_results"
            os.makedirs(output_dir, exist_ok=True)
            
            # 生成安全的文件名
            safe_filename = "".join(c for c in file_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_filename = f"OCR_{safe_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            doc.save(output_path)
            logger.info(f"OCR result saved to Word document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to save OCR result to Word: {e}")
            return ""

    def analyze_document(self, file_path: str, analysis_type: str) -> Dict[str, Any]:
        """完整的文档分析流程"""
        try:
            # 获取文件信息
            file_info = {
                "name": os.path.basename(file_path),
                "size": f"{os.path.getsize(file_path)} bytes"
            }
            
            # 第一步：OCR文本提取
            ocr_result = self.extract_text_with_ocr(file_path)
            
            if "error" in ocr_result:
                return {
                    "extracted_data": {
                        "error": f"OCR提取失败: {ocr_result['error']}"
                    },
                    "verification_results": {
                        "processed": "未通过"
                    },
                    "recommendations": [
                        "OCR文本提取失败",
                        "请检查文件格式和质量",
                        "建议重新上传清晰的文档"
                    ],
                    "confidence_score": 0.0
                }
            
            # 第二步：基于提取文本的智能分析
            extracted_text = ocr_result.get("text", "")
            
            # 保存OCR结果到Word文档
            word_path = ""
            if extracted_text.strip():
                word_path = self.save_extracted_text_to_word(extracted_text, file_info["name"])
            
            if not extracted_text.strip():
                return {
                    "extracted_data": {
                        "note": "未提取到文本内容",
                        "file_info": file_info,
                        "ocr_method": ocr_result.get("method", "Unknown")
                    },
                    "verification_results": {
                        "text_extraction": "失败"
                    },
                    "recommendations": [
                        "未能从文档中提取到文本",
                        "可能是图片质量问题或文档格式不支持",
                        "建议使用更清晰的扫描件或图片"
                    ],
                    "confidence_score": 0.0
                }
            
            # 执行智能分析
            analysis_result = self.analyze_document_with_ai(extracted_text, analysis_type, file_info)
            
            # 添加OCR相关信息
            analysis_result["ocr_info"] = {
                "method": ocr_result.get("method", "Unknown"),
                "pages": ocr_result.get("pages", 1),
                "text_length": len(extracted_text),
                "word_document": word_path
            }
            
            # 添加完整文本到结果中
            analysis_result["extracted_data"]["full_text"] = extracted_text
            
            logger.info(f"Document analysis completed for {file_info['name']}")
            return analysis_result
            
        except Exception as e:
            logger.error(f"Analysis failed: {str(e)}")
            return {
                "extracted_data": {
                    "error": f"分析过程出错: {str(e)}"
                },
                "verification_results": {
                    "processed": "出错"
                },
                "recommendations": [
                    "文档分析过程中出现错误",
                    "请检查文件完整性",
                    "如问题持续，请联系技术支持"
                ],
                "confidence_score": 0.0
            }

# 初始化分析器
try:
    analyzer = OptimizedDocumentAnalyzer()
    logger.info(f"Document analyzer initialized, OCR available: {analyzer.available}")
except Exception as e:
    logger.error(f"Failed to initialize analyzer: {e}")
    analyzer = None

@app.get("/")
async def root():
    """健康检查接口"""
    return {
        "status": "ok",
        "message": "Material Review Service with OCR is running",
        "ocr_available": analyzer.available if analyzer else False,
        "paddle_ocr_imported": PADDLE_OCR_AVAILABLE
    }

@app.post("/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...)
):
    """文档上传和OCR分析接口"""
    try:
        if not analyzer:
            raise HTTPException(status_code=500, detail="Document analyzer not initialized")
        
        logger.info(f"Received file upload: {file.filename}, type: {document_type}")
        
        # 检查文件类型
        allowed_extensions = {'.pdf', '.jpg', '.jpeg', '.png'}
        file_extension = os.path.splitext(file.filename.lower())[1]
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {file_extension}。支持的格式: PDF, JPG, PNG"
            )
        
        # 保存上传的文件
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)
        
        file_path = os.path.join(upload_dir, file.filename)
        
        # 写入文件
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        logger.info(f"File saved to: {file_path}")
        
        # OCR分析文档
        analysis_result = analyzer.analyze_document(file_path, document_type)
        
        # 清理临时文件
        try:
            os.remove(file_path)
        except:
            pass
        
        response_data = {
            "success": True,
            "message": "文档OCR分析完成",
            "analysis_result": analysis_result
        }
        
        # 如果生成了Word文档，添加下载链接
        if analysis_result.get("ocr_info", {}).get("word_document"):
            word_filename = os.path.basename(analysis_result["ocr_info"]["word_document"])
            response_data["word_download_url"] = f"/download-ocr-result/{word_filename}"
        
        return JSONResponse(content=response_data)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")

@app.get("/download-ocr-result/{filename}")
async def download_ocr_result(filename: str):
    """下载OCR结果Word文档"""
    try:
        file_path = os.path.join("ocr_results", filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 验证文件是否为.docx格式
        if not filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="不支持的文件格式")
        
        logger.info(f"Downloading OCR result: {filename}")
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件下载失败: {str(e)}")

if __name__ == "__main__":
    logger.info("Starting Material Review Service with OCR...")
    uvicorn.run(app, host="0.0.0.0", port=8003)