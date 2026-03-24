#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import logging
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import uvicorn
from typing import Dict, Any
import tempfile
import fitz  # PyMuPDF for PDF processing
from docx import Document
from datetime import datetime
import base64

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 尝试导入不同的OCR库
TESSERACT_AVAILABLE = False
PADDLE_OCR_AVAILABLE = False
EASYOCR_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
    logger.info("Tesseract OCR available")
except ImportError:
    logger.warning("Tesseract OCR not available")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
    logger.info("EasyOCR available")
except ImportError:
    logger.warning("EasyOCR not available")

try:
    from paddleocr import PaddleOCR
    PADDLE_OCR_AVAILABLE = True
    logger.info("PaddleOCR available")
except ImportError:
    logger.warning("PaddleOCR not available")

# 创建FastAPI应用
app = FastAPI(title="Alternative OCR Service", version="1.0.0")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3004"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class MultipleOCRAnalyzer:
    """使用多种OCR引擎的文档分析器"""
    
    def __init__(self):
        self.engines = {}
        self.available = False
        
        # 初始化Tesseract
        if TESSERACT_AVAILABLE:
            try:
                # 配置Tesseract（如果有的话）
                # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                self.engines['tesseract'] = True
                logger.info("Tesseract engine initialized")
            except Exception as e:
                logger.warning(f"Tesseract initialization failed: {e}")
        
        # 初始化EasyOCR
        if EASYOCR_AVAILABLE:
            try:
                self.engines['easyocr'] = easyocr.Reader(['en', 'ch_sim'])
                logger.info("EasyOCR engine initialized")
            except Exception as e:
                logger.warning(f"EasyOCR initialization failed: {e}")
        
        # 初始化PaddleOCR（最简配置）
        if PADDLE_OCR_AVAILABLE:
            try:
                self.engines['paddleocr'] = PaddleOCR(lang='en')  # 只用英文避免问题
                logger.info("PaddleOCR engine initialized")
            except Exception as e:
                logger.warning(f"PaddleOCR initialization failed: {e}")
        
        self.available = len(self.engines) > 0
        logger.info(f"Available OCR engines: {list(self.engines.keys())}")
    
    def pdf_to_images(self, pdf_path: str) -> list:
        """将PDF转换为高质量图片"""
        images = []
        try:
            pdf_document = fitz.open(pdf_path)
            logger.info(f"PDF has {len(pdf_document)} pages")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # 极高分辨率转换
                mat = fitz.Matrix(6.0, 6.0)  # 6倍缩放
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_data = pix.tobytes("png")
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                    tmp_file.write(img_data)
                    images.append(tmp_file.name)
                    logger.info(f"Created ultra high-res image for page {page_num + 1}")
            
            pdf_document.close()
            return images
        except Exception as e:
            logger.error(f"Failed to convert PDF: {e}")
            return []
    
    def extract_with_tesseract(self, image_path: str) -> str:
        """使用Tesseract提取文本"""
        try:
            img = Image.open(image_path)
            
            # 多种配置尝试
            configs = [
                '--psm 6',  # 单一文本块
                '--psm 4',  # 单列文本
                '--psm 3',  # 完全自动分割（默认）
                '--psm 11', # 稀疏文本
                '--psm 12', # 稀疏文本，OSD
            ]
            
            best_result = ""
            best_length = 0
            
            for config in configs:
                try:
                    text = pytesseract.image_to_string(img, config=config)
                    if len(text.strip()) > best_length:
                        best_result = text
                        best_length = len(text.strip())
                        logger.info(f"Tesseract config '{config}' extracted {best_length} chars")
                except Exception as e:
                    logger.warning(f"Tesseract config '{config}' failed: {e}")
            
            return best_result.strip()
        except Exception as e:
            logger.error(f"Tesseract extraction failed: {e}")
            return ""
    
    def extract_with_easyocr(self, image_path: str) -> str:
        """使用EasyOCR提取文本"""
        try:
            reader = self.engines['easyocr']
            results = reader.readtext(image_path, detail=0)
            text = '\n'.join(results)
            logger.info(f"EasyOCR extracted {len(text)} characters")
            return text
        except Exception as e:
            logger.error(f"EasyOCR extraction failed: {e}")
            return ""
    
    def extract_with_paddleocr(self, image_path: str) -> str:
        """使用PaddleOCR提取文本"""
        try:
            ocr = self.engines['paddleocr']
            result = ocr.predict(image_path)
            
            if result and len(result) > 0 and result[0]:
                text_lines = []
                for line in result[0]:
                    if len(line) >= 2:
                        text_content = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                        if text_content and text_content.strip():
                            text_lines.append(text_content.strip())
                
                text = '\n'.join(text_lines)
                logger.info(f"PaddleOCR extracted {len(text)} characters")
                return text
            return ""
        except Exception as e:
            logger.error(f"PaddleOCR extraction failed: {e}")
            return ""
    
    def extract_text_multiple_engines(self, file_path: str) -> Dict[str, Any]:
        """使用多个OCR引擎提取文本"""
        if not self.available:
            return {"error": "No OCR engines available", "text": ""}
        
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                image_paths = self.pdf_to_images(file_path)
                if not image_paths:
                    return {"error": "Failed to convert PDF", "text": ""}
                
                # 对每个引擎都尝试
                all_results = {}
                
                for page_idx, img_path in enumerate(image_paths):
                    page_results = {}
                    
                    # Tesseract
                    if 'tesseract' in self.engines:
                        tesseract_text = self.extract_with_tesseract(img_path)
                        if tesseract_text:
                            page_results['tesseract'] = tesseract_text
                    
                    # EasyOCR
                    if 'easyocr' in self.engines:
                        easyocr_text = self.extract_with_easyocr(img_path)
                        if easyocr_text:
                            page_results['easyocr'] = easyocr_text
                    
                    # PaddleOCR
                    if 'paddleocr' in self.engines:
                        paddleocr_text = self.extract_with_paddleocr(img_path)
                        if paddleocr_text:
                            page_results['paddleocr'] = paddleocr_text
                    
                    all_results[f'page_{page_idx + 1}'] = page_results
                    
                    # 清理临时文件
                    try:
                        os.unlink(img_path)
                    except:
                        pass
                
                # 选择最好的结果
                best_engine = None
                best_text = ""
                best_length = 0
                
                for page_data in all_results.values():
                    for engine, text in page_data.items():
                        if len(text) > best_length:
                            best_engine = engine
                            best_text = text
                            best_length = len(text)
                
                return {
                    "text": best_text,
                    "best_engine": best_engine,
                    "all_results": all_results,
                    "method": "PDF->Multiple_OCR"
                }
            
            else:
                # 直接处理图片
                results = {}
                
                if 'tesseract' in self.engines:
                    results['tesseract'] = self.extract_with_tesseract(file_path)
                
                if 'easyocr' in self.engines:
                    results['easyocr'] = self.extract_with_easyocr(file_path)
                
                if 'paddleocr' in self.engines:
                    results['paddleocr'] = self.extract_with_paddleocr(file_path)
                
                # 选择最好的结果
                best_engine = None
                best_text = ""
                best_length = 0
                
                for engine, text in results.items():
                    if len(text) > best_length:
                        best_engine = engine
                        best_text = text
                        best_length = len(text)
                
                return {
                    "text": best_text,
                    "best_engine": best_engine,
                    "all_results": results,
                    "method": "Direct_Multiple_OCR"
                }
        
        except Exception as e:
            logger.error(f"Multiple OCR extraction failed: {e}")
            return {"error": str(e), "text": ""}
    
    def save_ocr_comparison_to_word(self, ocr_results: Dict[str, Any], file_name: str) -> str:
        """保存OCR对比结果到Word文档"""
        try:
            doc = Document()
            
            # 标题
            doc.add_heading(f'多引擎OCR对比结果 - {file_name}', 0)
            
            # 基本信息
            doc.add_heading('分析信息', level=1)
            doc.add_paragraph(f'文件名: {file_name}')
            doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'最佳引擎: {ocr_results.get("best_engine", "未知")}')
            doc.add_paragraph(f'提取方法: {ocr_results.get("method", "未知")}')
            
            # 最佳结果
            doc.add_heading('最佳识别结果', level=1)
            best_text = ocr_results.get("text", "")
            if best_text:
                doc.add_paragraph(best_text)
            else:
                doc.add_paragraph('未提取到文本内容')
            
            # 所有引擎对比
            doc.add_heading('各引擎详细结果', level=1)
            all_results = ocr_results.get("all_results", {})
            
            for page_or_engine, content in all_results.items():
                doc.add_heading(f'{page_or_engine}', level=2)
                
                if isinstance(content, dict):
                    for engine, text in content.items():
                        doc.add_heading(f'{engine.upper()} 引擎', level=3)
                        doc.add_paragraph(f'字符数: {len(text)}')
                        if text:
                            doc.add_paragraph(text)
                        else:
                            doc.add_paragraph('无识别结果')
                        doc.add_paragraph('---')
                else:
                    doc.add_paragraph(str(content))
            
            # 保存文档
            output_dir = "ocr_results"
            os.makedirs(output_dir, exist_ok=True)
            
            safe_filename = "".join(c for c in file_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_filename = f"MultiOCR_{safe_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            doc.save(output_path)
            logger.info(f"OCR comparison saved: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to save OCR comparison: {e}")
            return ""

# 初始化分析器
try:
    analyzer = MultipleOCRAnalyzer()
    logger.info(f"Multiple OCR analyzer initialized, available: {analyzer.available}")
except Exception as e:
    logger.error(f"Failed to initialize analyzer: {e}")
    analyzer = None

@app.get("/")
async def root():
    """健康检查接口"""
    return {
        "status": "ok",
        "message": "Alternative OCR Service is running",
        "available_engines": list(analyzer.engines.keys()) if analyzer else [],
        "tesseract_available": TESSERACT_AVAILABLE,
        "easyocr_available": EASYOCR_AVAILABLE,
        "paddleocr_available": PADDLE_OCR_AVAILABLE
    }

@app.post("/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...)
):
    """文档上传和多引擎OCR分析接口"""
    try:
        if not analyzer or not analyzer.available:
            raise HTTPException(status_code=500, detail="No OCR engines available")
        
        logger.info(f"Received file: {file.filename}, type: {document_type}")
        
        # 检查文件类型
        allowed_extensions = {'.pdf', '.jpg', '.jpeg', '.png'}
        file_extension = os.path.splitext(file.filename.lower())[1]
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {file_extension}"
            )
        
        # 保存文件
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, file.filename)
        
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # 多引擎OCR分析
        ocr_results = analyzer.extract_text_multiple_engines(file_path)
        
        # 保存对比结果
        word_path = ""
        if ocr_results.get("text"):
            word_path = analyzer.save_ocr_comparison_to_word(ocr_results, file.filename)
        
        # 清理上传文件
        try:
            os.remove(file_path)
        except:
            pass
        
        # 构建分析结果
        extracted_text = ocr_results.get("text", "")
        analysis_result = {
            "extracted_data": {
                "document_type": document_type,
                "file_name": file.filename,
                "best_engine": ocr_results.get("best_engine", "未知"),
                "text_length": len(extracted_text),
                "full_text": extracted_text,
                "engines_tested": list(analyzer.engines.keys()),
                "extraction_method": ocr_results.get("method", "未知")
            },
            "verification_results": {
                "text_extracted": "通过" if extracted_text else "失败",
                "multiple_engines": "通过" if len(analyzer.engines) > 1 else "单引擎",
                "best_result_selected": "通过" if ocr_results.get("best_engine") else "失败"
            },
            "recommendations": [
                f"使用了 {len(analyzer.engines)} 个OCR引擎进行对比",
                f"最佳结果来自: {ocr_results.get('best_engine', '未知')}",
                f"提取了 {len(extracted_text)} 个字符" if extracted_text else "未提取到文本内容",
                "建议查看Word文档了解各引擎的详细对比结果"
            ],
            "confidence_score": 0.8 if extracted_text else 0.1,
            "ocr_info": {
                "method": ocr_results.get("method", "未知"),
                "word_document": word_path,
                "all_results": ocr_results.get("all_results", {})
            }
        }
        
        response_data = {
            "success": True,
            "message": "多引擎OCR分析完成",
            "analysis_result": analysis_result
        }
        
        if word_path:
            word_filename = os.path.basename(word_path)
            response_data["word_download_url"] = f"/download-ocr-result/{word_filename}"
        
        return JSONResponse(content=response_data)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")

@app.get("/download-ocr-result/{filename}")
async def download_ocr_result(filename: str):
    """下载OCR结果Word文档"""
    try:
        file_path = os.path.join("ocr_results", filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        if not filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="不支持的文件格式")
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件下载失败: {str(e)}")

if __name__ == "__main__":
    logger.info("Starting Alternative OCR Service...")
    uvicorn.run(app, host="0.0.0.0", port=8004)  # 使用不同端口避免冲突